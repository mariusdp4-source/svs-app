#!/usr/bin/env python3
"""SVS Stock App – Tornado bediener  (v2026-07-08)"""
import tornado.web, tornado.ioloop
import sqlite3, hashlib, json, os, datetime

DB_PATH = os.path.join(os.path.expanduser('~'), 'svs_stock.db')
PORT = 3000
SECRET_KEY = os.environ.get('SECRET_KEY', 'svs-stock-secret-2026')

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

# ─── Datum / Tyd Hulpfunksies ─────────────────────────────────────────────────

def sast_now():
    """Huidige datum-tyd in SAST (UTC+2) as 'YYYY-MM-DD HH:MM:SS'."""
    return (datetime.datetime.utcnow() + datetime.timedelta(hours=2)).strftime('%Y-%m-%d %H:%M:%S')

def week_monday(date_str=None):
    """Gee die ISO-datum van die Woensdag wat die voorraadweek begin vir 'n gegewe datum (of vandag in SAST).
    Voorraadweke loop Woensdag tot Dinsdag.
    Logika: keer terug na die mees onlangse Woensdag (gister of vroeër), selfs as dit 'n paar dae gelede was.
    So Ma 13 Jul → 8 Jul (Wo); Wo 15 Jul → 15 Jul; Do 16 Jul → 15 Jul."""
    if date_str:
        d = datetime.date.fromisoformat(str(date_str)[:10])
    else:
        d = (datetime.datetime.utcnow() + datetime.timedelta(hours=2)).date()
    # Dae sedert die mees onlangse Woensdag (0 as dit Woensdag is, 1 as Donderdag, ens.)
    days_since_wed = (d.weekday() - 2) % 7
    return (d - datetime.timedelta(days=days_since_wed)).isoformat()

def week_sunday(wednesday_str):
    """Gee die ISO-datum van die Dinsdag (laaste dag) van die voorraadweek wat op Woensdag begin.
    Gebruik as eindpunt vir BETWEEN-navrae: Woensdag + 6 dae = Dinsdag."""
    d = datetime.date.fromisoformat(str(wednesday_str)[:10])
    return (d + datetime.timedelta(days=6)).isoformat()

def week_friday(wednesday_str):
    """Gee die ISO-datum van die VOLGENDE Woensdag (vir vertoon as einde van die voorraadweek)."""
    d = datetime.date.fromisoformat(str(wednesday_str)[:10])
    return (d + datetime.timedelta(days=7)).isoformat()

MAANDE = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']
MONTHS_FULL = ['January','February','March','April','May','June','July','August','September','October','November','December']

def fmt_date(date_str):
    """Formateer 'YYYY-MM-DD' as '14 Jul 2026'."""
    try:
        d = datetime.date.fromisoformat(str(date_str)[:10])
        return f"{d.day} {MAANDE[d.month - 1]} {d.year}"
    except Exception:
        return str(date_str)

def order_window():
    """Returns (pack_date, is_open) for the current ordering window.
    Window: Tuesday 00:00 → Friday 11:00 SAST.
    Tuesday → pack_date is next Wednesday.
    Wed–Fri → pack_date is current Wednesday.
    """
    now = datetime.datetime.utcnow() + datetime.timedelta(hours=2)
    wd = now.weekday()  # 0=Mon,1=Tue,2=Wed,3=Thu,4=Fri,5=Sat,6=Sun
    if wd == 1:  # Tuesday → pack next Wednesday
        return (now.date() + datetime.timedelta(days=1)).isoformat(), True
    elif wd in (2, 3):  # Wed or Thu
        return week_monday(), True
    elif wd == 4 and now.hour < 11:  # Friday before 11:00
        return week_monday(), True
    else:
        return week_monday(), False

def check_password(stored, provided):
    try:
        salt, h = stored.split(':')
        test = hashlib.pbkdf2_hmac('sha256', provided.encode(), salt.encode(), 100000).hex()
        return test == h
    except Exception:
        return False

# ─── Base Handler ─────────────────────────────────────────────────────────────

class BaseHandler(tornado.web.RequestHandler):
    def get_current_user(self):
        uid = self.get_secure_cookie('svs_user')
        if not uid: return None
        try:
            db = get_db()
            u = db.execute("SELECT u.*, s.code as salon_code, s.name as salon_name FROM users u LEFT JOIN salons s ON u.salon_id = s.id WHERE u.id=?", (int(uid),)).fetchone()
            db.close()
            return dict(u) if u else None
        except Exception:
            return None

    def render_error(self, msg):
        self.write(f'<html><body style="font-family:Arial;padding:2rem"><h2 style="color:#c0392b">Fout</h2><p>{msg}</p><a href="/">Terug</a></body></html>')

# ─── Auth ─────────────────────────────────────────────────────────────────────

class LoginHandler(BaseHandler):
    def get(self):
        if self.current_user:
            self.redirect('/dashboard'); return
        self.render('login.html', error=None)

    def post(self):
        username = self.get_argument('username', '').strip()
        password = self.get_argument('password', '').strip()
        db = get_db()
        user = db.execute("SELECT * FROM users WHERE LOWER(username)=LOWER(?)", (username,)).fetchone()
        db.close()
        if user and check_password(user['password_hash'], password):
            self.set_secure_cookie('svs_user', str(user['id']), expires_days=7)
            self.redirect('/dashboard')
        else:
            self.render('login.html', error='Verkeerde gebruikersnaam of wagwoord.')

class LogoutHandler(BaseHandler):
    def get(self):
        self.clear_cookie('svs_user')
        self.redirect('/login')

class DashboardHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard')
        else:
            self.redirect('/salon/dashboard')

# ─── Salon ────────────────────────────────────────────────────────────────────

class SalonDashboardHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        db = get_db()
        upcoming_pack, window_open = order_window()
        draft = db.execute(
            "SELECT * FROM stock_orders WHERE salon_id=? AND status='draft'",
            (u['salon_id'],)
        ).fetchone()
        # Active order for the current/upcoming pack week (submitted or packed)
        current_week_order = db.execute("""
            SELECT * FROM stock_orders
            WHERE salon_id=? AND pack_date=? AND status IN ('submitted','packed')
            ORDER BY submitted_at DESC LIMIT 1
        """, (u['salon_id'], upcoming_pack)).fetchone()
        current_week_order = dict(current_week_order) if current_week_order else None
        active_orders_raw = db.execute(
            "SELECT * FROM stock_orders WHERE salon_id=? AND status IN ('submitted','packed') ORDER BY pack_date DESC",
            (u['salon_id'],)
        ).fetchall()
        # Bepaal watter gepakte bestellings uitstaande items het
        packed_active_ids = [o['id'] for o in active_orders_raw if o['status'] == 'packed']
        salon_outstanding_ids = set()
        if packed_active_ids:
            ph2 = ','.join('?' * len(packed_active_ids))
            out2 = db.execute(f"""
                SELECT oi.order_id FROM order_items oi
                WHERE oi.order_id IN ({ph2})
                  AND oi.category != 'Handoeke'
                  AND oi.quantity > CASE WHEN COALESCE(oi.packed,0)=1
                                         THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                                         ELSE 0 END
                GROUP BY oi.order_id
            """, packed_active_ids).fetchall()
            salon_outstanding_ids = {r['order_id'] for r in out2}
        active_orders = []
        for o in active_orders_raw:
            od = dict(o)
            od['has_outstanding'] = (o['id'] in salon_outstanding_ids)
            active_orders.append(od)
        delivered_orders = db.execute(
            "SELECT * FROM stock_orders WHERE salon_id=? AND status='delivered' ORDER BY delivered_at DESC LIMIT 20",
            (u['salon_id'],)
        ).fetchall()
        # Uitstaande items van laaste aflewering
        last_delivery = db.execute("""
            SELECT id, pack_date, delivered_at FROM stock_orders
            WHERE salon_id=? AND status='delivered'
            ORDER BY delivered_at DESC LIMIT 1
        """, (u['salon_id'],)).fetchone()
        outstanding_last = []
        if last_delivery:
            rows = db.execute("""
                SELECT id, product_name, category,
                       quantity as ordered_qty,
                       COALESCE(delivered_qty, 0) as del_qty,
                       quantity - COALESCE(delivered_qty, 0) as outstanding_qty
                FROM order_items
                WHERE order_id=? AND quantity > COALESCE(delivered_qty, 0)
                  AND category != 'Handoeke'
                ORDER BY category, product_name
            """, (last_delivery['id'],)).fetchall()
            outstanding_last = [dict(r) for r in rows]
        # Meeste onlangse handoeke-rekords (uit aflewering OF staan-alleen)
        last_towel_logs = []
        latest_towel = db.execute("""
            SELECT date(created_at) as log_date
            FROM towel_logs WHERE salon_id=?
            ORDER BY created_at DESC LIMIT 1
        """, (u['salon_id'],)).fetchone()
        if latest_towel:
            last_towel_logs = [dict(r) for r in db.execute("""
                SELECT towel_type, sent, collected FROM towel_logs
                WHERE salon_id=? AND date(created_at) = ?
                  AND (sent > 0 OR collected > 0)
                ORDER BY towel_type
            """, (u['salon_id'], latest_towel['log_date'])).fetchall()]
        # Aktiewe bestelling se handoeke (Opgetel Vuil = wat die salon bestel het)
        active_order_towels = {}
        active_order_q = db.execute("""
            SELECT id FROM stock_orders
            WHERE salon_id=? AND status IN ('submitted','packed')
            ORDER BY submitted_at DESC LIMIT 1
        """, (u['salon_id'],)).fetchone()
        if active_order_q:
            for r in db.execute("""
                SELECT product_name, quantity FROM order_items
                WHERE order_id=? AND category='Handoeke' AND quantity > 0
            """, (active_order_q['id'],)).fetchall():
                active_order_towels[r['product_name']] = float(r['quantity'])
        # Mees onlangse skoon handoeke gestuur per tipe (Terug Ontvang Skoon)
        last_towel_sent = {}
        for r in db.execute("""
            SELECT towel_type, sent FROM towel_logs
            WHERE salon_id=? AND sent > 0
            ORDER BY created_at DESC
        """, (u['salon_id'],)).fetchall():
            if r['towel_type'] not in last_towel_sent:
                last_towel_sent[r['towel_type']] = float(r['sent'])
        db.close()
        self.render('salon/dashboard.html', user=u, draft=draft,
                    active_orders=active_orders, delivered_orders=delivered_orders,
                    outstanding_last=outstanding_last,
                    last_delivery=dict(last_delivery) if last_delivery else None,
                    last_towel_logs=last_towel_logs,
                    window_open=window_open, current_week_order=current_week_order,
                    active_order_towels=active_order_towels,
                    last_towel_sent=last_towel_sent,
                    towel_types=TOWEL_TYPES)

class OrderNewHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        upcoming_pack, window_open = order_window()
        if not window_open:
            self.redirect('/salon/dashboard'); return
        db = get_db()
        # Block second order: if submitted/packed order already exists for this week, redirect
        existing = db.execute("""
            SELECT id FROM stock_orders
            WHERE salon_id=? AND pack_date=? AND status IN ('submitted','packed')
        """, (u['salon_id'], upcoming_pack)).fetchone()
        if existing:
            db.close(); self.redirect('/salon/dashboard'); return
        draft = db.execute(
            "SELECT * FROM stock_orders WHERE salon_id=? AND status='draft'",
            (u['salon_id'],)
        ).fetchone()
        if draft:
            order_id = draft['id']
        else:
            cur = db.execute(
                "INSERT INTO stock_orders (salon_id, pack_date, status) VALUES (?,?,'draft')",
                (u['salon_id'], upcoming_pack)
            )
            db.commit()
            order_id = cur.lastrowid
        # Laai bestaande konsep items
        draft_qtys = {}
        if draft:
            rows = db.execute("SELECT product_name, quantity FROM order_items WHERE order_id=?", (draft['id'],)).fetchall()
            for r in rows:
                draft_qtys[r['product_name']] = r['quantity']

        # Dra uitstaande items van laaste aflewering oor (slegs vir nuwe bestellings)
        carried_forward = []
        if not draft:
            last_order = db.execute("""
                SELECT id FROM stock_orders
                WHERE salon_id=? AND status='delivered'
                ORDER BY delivered_at DESC LIMIT 1
            """, (u['salon_id'],)).fetchone()
            if last_order:
                cf_rows = db.execute("""
                    SELECT product_name, category,
                           quantity - COALESCE(delivered_qty, 0) as outstanding_qty
                    FROM order_items
                    WHERE order_id=? AND category != 'Handoeke'
                      AND quantity > COALESCE(delivered_qty, 0)
                    ORDER BY category, product_name
                """, (last_order['id'],)).fetchall()
                for r in cf_rows:
                    carried_forward.append({
                        'product_name':  r['product_name'],
                        'category':      r['category'],
                        'outstanding_qty': float(r['outstanding_qty'])
                    })
                    draft_qtys[r['product_name']] = float(r['outstanding_qty'])

        salon_products = db.execute(
            "SELECT * FROM products WHERE category='Salon' AND active=1 ORDER BY name COLLATE NOCASE"
        ).fetchall()
        cleaning_products = db.execute(
            "SELECT * FROM products WHERE category='Cleaning' AND active=1 ORDER BY name COLLATE NOCASE"
        ).fetchall()
        perm_products = db.execute(
            "SELECT * FROM products WHERE category='Perms' AND active=1 ORDER BY name COLLATE NOCASE"
        ).fetchall()
        retail_products = db.execute(
            "SELECT * FROM products WHERE category='Retail' AND active=1 ORDER BY name COLLATE NOCASE"
        ).fetchall()
        tint_products = db.execute(
            "SELECT * FROM products WHERE category='Tints' AND active=1 ORDER BY name COLLATE NOCASE"
        ).fetchall()

        # Laai meeste onlangse handoeke-data (aflewering OF staan-alleen, mees onlangse per tipe)
        last_towels = {}  # {towel_type: {sent, collected, next_qty}}
        for r in db.execute("""
            SELECT towel_type, sent, collected, COALESCE(next_qty, 0) as next_qty
            FROM towel_logs WHERE salon_id=? ORDER BY created_at DESC
        """, (u['salon_id'],)).fetchall():
            if r['towel_type'] not in last_towels:
                last_towels[r['towel_type']] = {
                    'sent':      float(r['sent'] or 0),
                    'collected': float(r['collected'] or 0),
                    'next_qty':  float(r['next_qty'] or 0)
                }

        # Laai huidige konsep se handoek-hoeveelhede en notas (as dit bestaan)
        draft_towel_qtys = {}
        draft_towel_notes = {}
        draft_custom_towels = []   # towel types not in TOWEL_TYPES (user-added)
        if draft:
            for r in db.execute(
                "SELECT product_name, quantity, note FROM order_items WHERE order_id=? AND category='Handoeke'",
                (draft['id'],)
            ).fetchall():
                draft_towel_qtys[r['product_name']] = float(r['quantity'])
                draft_towel_notes[r['product_name']] = r['note'] or ''
                if r['product_name'] not in TOWEL_TYPES:
                    draft_custom_towels.append({
                        'name': r['product_name'],
                        'qty':  int(float(r['quantity'])),
                        'note': r['note'] or ''
                    })

        # Laai tints van huidige konsep (sodat hulle nie verlore gaan as bestelling teruggekeer word)
        draft_tints = []
        if draft:
            tint_rows = db.execute(
                "SELECT product_name, quantity FROM order_items WHERE order_id=? AND category='Tints' ORDER BY id",
                (draft['id'],)
            ).fetchall()
            draft_tints = [{'name': r['product_name'], 'qty': int(float(r['quantity']))} for r in tint_rows]

        # Laai aangepaste produk-items van konsep (Salon/Cleaning/Perms/Retail, nie Tints/Handoeke)
        # sodat hulle weer vertoon word as die gebruiker teruggaan na "← Add Items"
        draft_custom_products = {'Salon': [], 'Cleaning': [], 'Perms': [], 'Retail': []}
        if draft:
            for r in db.execute(
                "SELECT product_name, category, quantity, note FROM order_items "
                "WHERE order_id=? AND is_custom=1 AND category NOT IN ('Tints','Handoeke') "
                "ORDER BY id",
                (draft['id'],)
            ).fetchall():
                cat = r['category']
                if cat in draft_custom_products:
                    draft_custom_products[cat].append({
                        'name': r['product_name'],
                        'qty':  int(float(r['quantity'])),
                        'note': r['note'] or ''
                    })

        db.close()
        today = datetime.date.today().strftime('%Y-%m-%d')
        self.render('salon/order_new.html', user=u, order_id=order_id,
                    salon_products=salon_products, cleaning_products=cleaning_products,
                    perm_products=perm_products, retail_products=retail_products,
                    tint_products=tint_products,
                    today=today, draft_qtys=draft_qtys, carried_forward=carried_forward,
                    towel_types=TOWEL_TYPES, last_towels=last_towels,
                    draft_towel_qtys=draft_towel_qtys, draft_towel_notes=draft_towel_notes,
                    draft_custom_towels=draft_custom_towels,
                    draft_custom_products=draft_custom_products,
                    form_action='/salon/order/submit',
                    back_url='/salon/dashboard',
                    submit_label='View Order',
                    ho_edit=False, ho_salon_name='')

class OrderSubmitHandler(BaseHandler):
    @tornado.web.authenticated
    def post(self):
        u = self.current_user
        order_id   = int(self.get_argument('order_id'))
        # Normaliseer altyd na die Woensdag van die gekose week
        pack_date  = week_monday(self.get_argument('pack_date', None) or datetime.date.today().isoformat())
        notes      = self.get_argument('notes', '')
        try:
            items = json.loads(self.get_argument('items_json', '[]'))
        except Exception:
            items = []

        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND salon_id=?",
            (order_id, u['salon_id'])
        ).fetchone()
        if not order:
            db.close(); self.render_error("Bestelling nie gevind nie"); return

        # Save as draft (not yet submitted) – user still needs to confirm on preview page
        db.execute(
            "UPDATE stock_orders SET pack_date=?, notes=?, status='draft' WHERE id=?",
            (pack_date, notes, order_id)
        )
        db.execute("DELETE FROM order_items WHERE order_id=?", (order_id,))
        for item in items:
            if not item.get('name', '').strip():
                continue
            # Moenie 0-hoeveelheid Handoeke stoor nie (voorkom vals "draft" wees)
            if item.get('category') == 'Handoeke' and float(item.get('quantity', 0)) == 0:
                continue
            db.execute(
                "INSERT INTO order_items (order_id, product_name, category, quantity, note, is_custom) VALUES (?,?,?,?,?,?)",
                (order_id, item['name'], item.get('category','Salon'),
                 float(item.get('quantity',1)), item.get('note',''),
                 1 if item.get('is_custom') else 0)
            )
        db.commit()
        db.close()
        # Redirect to preview page so user can review/edit before final submission
        self.redirect(f'/salon/order/preview/{order_id}')

class OrderViewHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self, order_id):
        u = self.current_user
        db = get_db()
        order = db.execute("SELECT * FROM stock_orders WHERE id=?", (order_id,)).fetchone()
        if not order:
            db.close(); self.render_error("Bestelling nie gevind nie"); return
        if u['role'] != 'ho_admin' and order['salon_id'] != u['salon_id']:
            db.close(); self.render_error("Geen toegang"); return
        all_items = db.execute(
            "SELECT * FROM order_items WHERE order_id=? ORDER BY category, product_name",
            (order_id,)
        ).fetchall()
        items       = [dict(i) for i in all_items if i['category'] != 'Handoeke']
        towel_items = [dict(i) for i in all_items if i['category'] == 'Handoeke']
        # Werklike aflewerings-towel data (slegs vir afgelewerde bestellings)
        towel_logs_data = []
        if dict(order)['status'] == 'delivered':
            towel_logs_data = [dict(r) for r in db.execute(
                "SELECT * FROM towel_logs WHERE order_id=? ORDER BY towel_type",
                (order_id,)
            ).fetchall()]
        salon = db.execute("SELECT * FROM salons WHERE id=?", (order['salon_id'],)).fetchone()
        db.close()
        # Packing progress — quantity-based (same logic as HO order view)
        total_qty  = sum(float(i['quantity']) for i in all_items if i['category'] != 'Handoeke')
        packed_qty = sum(float(i['packed_qty'] or 0) for i in all_items
                         if i['packed'] and i['category'] != 'Handoeke')
        self.render('salon/order_view.html', user=u, order=dict(order),
                    items=items, towel_items=towel_items,
                    towel_logs_data=towel_logs_data, salon=dict(salon),
                    total_qty=total_qty, packed_qty=packed_qty)

# ─── HO ───────────────────────────────────────────────────────────────────────

class HODashboardHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        # Groepeer bestellings per week (gebruik Maandag as sleutel)
        orders = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.status != 'draft'
            ORDER BY o.pack_date DESC, s.code
        """).fetchall()
        # Kry unieke weke as Maandag-datums (nuutste eerste)
        weeks = []
        seen = set()
        for o in orders:
            mon = week_monday(o['pack_date'])
            if mon not in seen:
                seen.add(mon)
                fri = week_friday(mon)
                weeks.append({
                    'monday':    mon,
                    'friday':    fri,
                    'label':     f"{fmt_date(mon)} – {fmt_date(fri)}",
                })
        # Huidige week se salon-status vir dashboard-kaarte
        cur_monday = week_monday()
        cur_sunday = week_sunday(cur_monday)
        cur_friday = week_friday(cur_monday)
        all_salons = db.execute(
            "SELECT id, code, name FROM salons WHERE active=1 ORDER BY code"
        ).fetchall()
        cur_week_orders = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.pack_date BETWEEN ? AND ? AND o.status != 'draft'
        """, (cur_monday, cur_sunday)).fetchall()
        cur_week_map = {o['salon_code']: dict(o) for o in cur_week_orders}
        # Bereken has_outstanding dinamies — soos HOWeekHandler dit doen
        outstanding_order_ids = set()
        packed_ids = [o['id'] for o in cur_week_orders if o['status'] == 'packed']
        if packed_ids:
            ph = ','.join('?' * len(packed_ids))
            out_rows = db.execute(f"""
                SELECT oi.order_id
                FROM order_items oi
                WHERE oi.order_id IN ({ph})
                  AND oi.category != 'Handoeke'
                  AND oi.quantity > CASE WHEN COALESCE(oi.packed,0)=1
                                         THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                                         ELSE 0 END
                GROUP BY oi.order_id
            """, packed_ids).fetchall()
            outstanding_order_ids = {r['order_id'] for r in out_rows}
        cur_week_salons = []
        for s in all_salons:
            order = cur_week_map.get(s['code'])
            cur_week_salons.append({
                'code':         s['code'],
                'name':         s['name'],
                'status':       order['status'] if order else None,
                'submitted_at': order['submitted_at'] if order else None,
                'order_id':     order['id'] if order else None,
                'has_outstanding': (order['id'] in outstanding_order_ids) if order else False,
            })
        # Wag op indiening (submitted maar nie gepak nie)
        pending = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.status = 'submitted'
            ORDER BY o.submitted_at DESC
        """).fetchall()
        db.close()
        self.render('ho/dashboard.html', user=u, weeks=weeks, pending=pending,
                    cur_week_salons=cur_week_salons,
                    cur_monday=cur_monday, cur_friday=cur_friday)

class HOWeekHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        # Normaliseer pack_date na die Maandag van die week
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        friday = week_friday(monday)
        db = get_db()
        # Fetch all non-draft orders; deduplicate per salon — prefer submitted/packed over delivered
        orders_raw = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.pack_date BETWEEN ? AND ? AND o.status != 'draft'
            ORDER BY s.code,
                     CASE o.status WHEN 'submitted' THEN 0 WHEN 'packed' THEN 0 ELSE 1 END,
                     o.submitted_at DESC
        """, (monday, sunday)).fetchall()
        seen_codes = set()
        orders = []
        for o in orders_raw:
            od = dict(o)
            if od['salon_code'] not in seen_codes:
                seen_codes.add(od['salon_code'])
                orders.append(od)

        # Gesamentlike item-lys: alleen vir die gekose orders
        selected_ids = [o['id'] for o in orders]
        if selected_ids:
            placeholders = ','.join(['?' for _ in selected_ids])
            items_raw = db.execute(f"""
                SELECT oi.order_id, oi.product_name, oi.category, oi.quantity,
                       COALESCE(oi.packed, 0) as packed,
                       CASE WHEN COALESCE(oi.packed,0)=1
                            THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                            ELSE 0 END as item_packed_qty,
                       s.code as salon_code
                FROM order_items oi
                JOIN stock_orders so ON oi.order_id = so.id
                JOIN salons s ON so.salon_id = s.id
                WHERE oi.order_id IN ({placeholders})
                ORDER BY oi.category, oi.product_name
            """, selected_ids).fetchall()
        else:
            items_raw = []
        db.close()

        # Track whether each order has any outstanding (not fully packed) items
        order_outstanding = {}   # order_id -> bool
        for row in items_raw:
            if row['category'] == 'Handoeke':
                continue
            oid = row['order_id']
            if row['quantity'] > row['item_packed_qty']:
                order_outstanding[oid] = True
            elif oid not in order_outstanding:
                order_outstanding[oid] = False

        for o in orders:
            o['has_outstanding'] = order_outstanding.get(o['id'], True)

        # Build combined list — merge duplicate salon entries per product
        from collections import OrderedDict
        agg = OrderedDict()
        for row in items_raw:
            key = (row['category'], row['product_name'])
            if key not in agg:
                agg[key] = {
                    'category':     row['category'],
                    'product_name': row['product_name'],
                    'total_qty':    0,
                    'packed_qty':   0,
                    'packed_salons': 0,
                    'total_salons': 0,
                    'salons': []
                }
            qty    = row['quantity']
            pq     = row['item_packed_qty']
            agg[key]['total_qty']  += qty
            agg[key]['packed_qty'] += pq
            # Merge by salon code so same salon only appears once per product
            existing = next((s for s in agg[key]['salons'] if s['code'] == row['salon_code']), None)
            if existing:
                existing['qty']        += qty
                existing['packed_qty'] += pq
                if not row['packed']:
                    existing['packed'] = False
            else:
                agg[key]['salons'].append({
                    'code':       row['salon_code'],
                    'qty':        qty,
                    'packed_qty': pq,
                    'packed':     bool(row['packed'])
                })
                agg[key]['total_salons'] += 1
                if row['packed']:
                    agg[key]['packed_salons'] += 1

        all_items     = list(agg.values())
        product_items = [i for i in all_items if i['category'] != 'Handoeke']
        towel_items   = sorted([i for i in all_items if i['category'] == 'Handoeke'],
                               key=lambda x: x['product_name'])
        # Outstanding = any product where actual packed qty < total ordered qty
        outstanding = [i for i in product_items
                       if i['packed_qty'] < i['total_qty']]

        # Towel matrix for the Towels tab
        SALON_ORDER_T = ['AV','BB','BV','EL','HE','PO','RS','TR','VG','WK']
        towel_salons  = [c for c in SALON_ORDER_T
                         if any(s['code'] == c for t in towel_items for s in t['salons'])]
        towel_matrix  = {t['product_name']: {s['code']: s['qty'] for s in t['salons']}
                         for t in towel_items}

        # Laai verskaffer-bestellings vir hierdie week
        db2 = get_db()
        sup_rows = db2.execute(
            "SELECT * FROM supplier_orders WHERE pack_date=?", (monday,)
        ).fetchall()
        db2.close()
        sup = {r['product_name']: dict(r) for r in sup_rows}
        for item in outstanding:
            rec = sup.get(item['product_name'], {})
            item['qty_needed']  = item['total_qty'] - item['packed_qty']
            item['ordered']     = rec.get('ordered', 0)
            item['qty_ordered'] = rec.get('qty_ordered', item['qty_needed'])
            item['notes']       = rec.get('notes', '')

        # Build supplier groups for Supplier Order tab
        db3 = get_db()
        prod_rows = db3.execute(
            "SELECT name, COALESCE(supplier,'') as supplier, category FROM products WHERE active=1 ORDER BY name"
        ).fetchall()
        db3.close()
        prod_sup_map  = {r['name'].lower().strip(): (r['supplier'] or ('HHB' if r['category'] == 'Tints' else 'No Supplier')) for r in prod_rows}
        all_products  = [r['name'] for r in prod_rows]
        tint_products = [r['name'] for r in prod_rows if r['category'] == 'Tints']

        supplier_groups = OrderedDict()
        sup_idx = 0
        for item in outstanding:
            _fallback = 'HHB' if item.get('category') == 'Tints' else 'No Supplier'
            sup_key = prod_sup_map.get(item['product_name'].lower().strip()) or _fallback
            item['_sup_idx'] = sup_idx
            sup_idx += 1
            if sup_key not in supplier_groups:
                supplier_groups[sup_key] = []
            supplier_groups[sup_key].append(item)

        # Also add extra items from supplier_orders not in current outstanding
        outstanding_names_lower = {item['product_name'].lower().strip() for item in outstanding}
        for pname, rec in sup.items():
            if pname.lower().strip() not in outstanding_names_lower and float(rec.get('qty_ordered') or 0) > 0:
                _cat_fb  = rec.get('category', 'Salon')
                _sup_fb  = 'HHB' if _cat_fb == 'Tints' else 'No Supplier'
                sup_key  = prod_sup_map.get(pname.lower().strip()) or _sup_fb
                extra = {
                    'product_name': pname,
                    'category':     rec.get('category', 'Salon'),
                    'qty_needed':   0,
                    'qty_ordered':  float(rec.get('qty_ordered') or 0),
                    'notes':        rec.get('notes', ''),
                    'ordered':      1,
                    '_sup_idx':     sup_idx,
                }
                sup_idx += 1
                if sup_key not in supplier_groups:
                    supplier_groups[sup_key] = []
                supplier_groups[sup_key].append(extra)

        total_sup_idx = sup_idx  # Starting index for JS-added extra rows

        self.render('ho/week.html', user=u, pack_date=monday,
                    week_friday=friday,
                    orders=orders, all_items=all_items,
                    product_items=product_items,
                    towel_items=towel_items,
                    towel_salons=towel_salons,
                    towel_matrix=towel_matrix,
                    outstanding=outstanding,
                    supplier_groups=supplier_groups,
                    all_products=all_products,
                    tint_products=tint_products,
                    total_sup_idx=total_sup_idx)

class HOOrderViewHandler(BaseHandler):
    """HO-spesifieke bestel-pakbladsy met per-item aftik."""
    @tornado.web.authenticated
    def get(self, order_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        order = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.id=? AND o.status != 'draft'
        """, (order_id,)).fetchone()
        if not order:
            db.close(); self.redirect('/ho/dashboard'); return
        items = db.execute(
            "SELECT * FROM order_items WHERE order_id=? ORDER BY category, product_name",
            (order_id,)
        ).fetchall()
        db.close()
        total  = sum(float(i['quantity']) for i in items if i['category'] != 'Handoeke')
        packed = sum(float(i['packed_qty'] or 0) for i in items
                     if i['packed'] and i['category'] != 'Handoeke')
        self.render('ho/order.html', user=u,
                    order=dict(order), items=[dict(i) for i in items],
                    total=total, packed_count=packed)

class HOOrderPackItemsHandler(BaseHandler):
    """Stoor per-item pakstatus; merk bestelling as 'gepak' as alles afgemerk is."""
    @tornado.web.authenticated
    def post(self, order_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND status != 'draft'", (order_id,)
        ).fetchone()
        if not order:
            db.close(); self.redirect('/ho/dashboard'); return
        # Delivered orders are locked — no further changes allowed
        if order['status'] == 'delivered':
            db.close()
            self.redirect(f'/ho/order/{order_id}')
            return
        items = db.execute("SELECT id, quantity FROM order_items WHERE order_id=?", (order_id,)).fetchall()
        any_packed = False
        for item in items:
            iid      = item['id']
            ord_qty  = item['quantity']
            is_packed = 1 if self.get_argument(f'pack_{iid}', '') == 'on' else 0
            try:
                pqty = float(self.get_argument(f'pqty_{iid}', ''))
                if pqty < 0:
                    pqty = 0
            except (ValueError, TypeError):
                pqty = ord_qty if is_packed else 0
            if not is_packed:
                pqty = 0
            elif pqty == 0:
                pqty = ord_qty   # default: full quantity if packed but 0 entered
            if is_packed:
                any_packed = True
            db.execute("UPDATE order_items SET packed=?, packed_qty=? WHERE id=?",
                       (is_packed, pqty, iid))
        # Sluit bestelling sodra ENIGSTE item gepak is — salon kan nie meer verander nie
        if any_packed:
            db.execute(
                "UPDATE stock_orders SET status='packed', packed_at=COALESCE(packed_at, ?) WHERE id=?",
                (sast_now(), order_id)
            )
        else:
            # Niks gepak nie — sit terug na submitted (salon kan dan nog verander)
            db.execute(
                "UPDATE stock_orders SET status='submitted', packed_at=NULL WHERE id=?",
                (order_id,)
            )
        db.commit()
        pack_date = week_monday(order['pack_date'])
        db.close()
        self.redirect(f'/ho/week/{pack_date}')

class HOFulfillOutstandingHandler(BaseHandler):
    """Merk alle uitstaande items in 'n afgelewerde bestelling as volledig afgelewer."""
    @tornado.web.authenticated
    def post(self, order_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND status='delivered'", (order_id,)
        ).fetchone()
        if order:
            # Only mark items that HO has actually packed (packed=1)
            # Items not yet packed remain outstanding
            db.execute("""
                UPDATE order_items
                SET delivered_qty = COALESCE(NULLIF(packed_qty,0), quantity)
                WHERE order_id=? AND category != 'Handoeke'
                  AND COALESCE(packed,0) = 1
                  AND COALESCE(delivered_qty, 0) < quantity
            """, (order_id,))
            db.commit()
        db.close()
        self.redirect(f'/ho/order/{order_id}')


class HOOrderEditHandler(BaseHandler):
    """Laat HO toe om 'n bestelling (draft/submitted/packed) direk te redigeer."""
    @tornado.web.authenticated
    def get(self, order_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        order = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.id=? AND o.status IN ('draft','submitted','packed')
        """, (order_id,)).fetchone()
        if not order:
            db.close()
            self.redirect('/ho/dashboard'); return

        # Produkte
        salon_products   = db.execute("SELECT * FROM products WHERE category='Salon'    AND active=1 ORDER BY name COLLATE NOCASE").fetchall()
        cleaning_products= db.execute("SELECT * FROM products WHERE category='Cleaning' AND active=1 ORDER BY name COLLATE NOCASE").fetchall()
        perm_products    = db.execute("SELECT * FROM products WHERE category='Perms'    AND active=1 ORDER BY name COLLATE NOCASE").fetchall()
        retail_products  = db.execute("SELECT * FROM products WHERE category='Retail'   AND active=1 ORDER BY name COLLATE NOCASE").fetchall()
        tint_products    = db.execute("SELECT * FROM products WHERE category='Tints'    AND active=1 ORDER BY name COLLATE NOCASE").fetchall()

        # Bestaande konsep items
        draft_qtys         = {}
        draft_towel_qtys   = {}
        draft_towel_notes  = {}
        draft_tints        = []
        draft_custom_towels= []
        all_items = db.execute(
            "SELECT * FROM order_items WHERE order_id=? ORDER BY category, product_name",
            (order_id,)
        ).fetchall()
        draft_custom_products = {'Salon': [], 'Cleaning': [], 'Perms': [], 'Retail': []}
        for r in all_items:
            if r['category'] == 'Handoeke':
                draft_towel_qtys[r['product_name']] = float(r['quantity'])
                draft_towel_notes[r['product_name']] = r['note'] or ''
                if r['product_name'] not in TOWEL_TYPES:
                    draft_custom_towels.append({
                        'name': r['product_name'],
                        'qty':  int(float(r['quantity'])),
                        'note': r['note'] or ''
                    })
            else:
                draft_qtys[r['product_name']] = r['quantity']
                if r['is_custom'] and r['category'] in draft_custom_products:
                    draft_custom_products[r['category']].append({
                        'name': r['product_name'],
                        'qty':  int(float(r['quantity'])),
                        'note': r['note'] or ''
                    })

        # Handoek-geskiedenis vir hierdie salon
        last_towels = {}
        for r in db.execute("""
            SELECT towel_type, sent, collected, COALESCE(next_qty, 0) as next_qty
            FROM towel_logs WHERE salon_id=? ORDER BY created_at DESC
        """, (order['salon_id'],)).fetchall():
            if r['towel_type'] not in last_towels:
                last_towels[r['towel_type']] = {
                    'sent':      float(r['sent'] or 0),
                    'collected': float(r['collected'] or 0),
                    'next_qty':  float(r['next_qty'] or 0)
                }

        db.close()
        today = dict(order).get('pack_date') or datetime.date.today().isoformat()
        self.render('salon/order_new.html', user=u, order_id=int(order_id),
                    salon_products=salon_products, cleaning_products=cleaning_products,
                    perm_products=perm_products, retail_products=retail_products,
                    tint_products=tint_products,
                    today=today, draft_qtys=draft_qtys, carried_forward=[],
                    towel_types=TOWEL_TYPES, last_towels=last_towels,
                    draft_towel_qtys=draft_towel_qtys, draft_towel_notes=draft_towel_notes,
                    draft_custom_towels=draft_custom_towels,
                    draft_custom_products=draft_custom_products,
                    # HO-spesifieke oorskrywings
                    form_action=f'/ho/order/{order_id}/edit',
                    back_url=f'/ho/dashboard',
                    submit_label='Save & Resubmit',
                    ho_edit=True,
                    ho_salon_name=f"{dict(order)['salon_code']} – {dict(order)['salon_name']}")

    @tornado.web.authenticated
    def post(self, order_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        pack_date = week_monday(self.get_argument('pack_date', None) or datetime.date.today().isoformat())
        notes     = self.get_argument('notes', '')
        try:
            items = json.loads(self.get_argument('items_json', '[]'))
        except Exception:
            items = []

        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND status IN ('draft','submitted','packed')", (order_id,)
        ).fetchone()
        if not order:
            db.close(); self.redirect('/ho/dashboard'); return

        current_status = order['status']

        # Slim items merge — bewaar packing data vir bestaande items
        existing_rows = db.execute(
            "SELECT * FROM order_items WHERE order_id=?", (order_id,)
        ).fetchall()
        existing = {(r['product_name'], r['category']): dict(r) for r in existing_rows}

        keep_keys = set()
        for item in items:
            name = item.get('name', '').strip()
            cat  = item.get('category', 'Salon')
            if not name:
                continue
            if cat == 'Handoeke' and float(item.get('quantity', 0)) == 0:
                continue
            key = (name, cat)
            keep_keys.add(key)
            qty  = float(item.get('quantity', 1))
            note = item.get('note', '')
            if key in existing:
                db.execute(
                    "UPDATE order_items SET quantity=?, note=? WHERE id=?",
                    (qty, note, existing[key]['id'])
                )
            else:
                db.execute(
                    "INSERT INTO order_items (order_id, product_name, category, quantity, note, is_custom, ho_added) VALUES (?,?,?,?,?,?,1)",
                    (order_id, name, cat, qty, note, 1 if item.get('is_custom') else 0)
                )

        # Verwyder items wat uitgehaal is
        for (name, cat), row in existing.items():
            if (name, cat) not in keep_keys:
                db.execute("DELETE FROM order_items WHERE id=?", (row['id'],))

        # Hou status onveranderd (submitted/packed bly soos is); draft word submitted
        if current_status == 'draft':
            db.execute(
                "UPDATE stock_orders SET pack_date=?, notes=?, status='submitted', submitted_at=? WHERE id=?",
                (week_monday(), notes, sast_now(), order_id)
            )
        else:
            db.execute(
                "UPDATE stock_orders SET notes=? WHERE id=?",
                (notes, order_id)
            )
        db.commit()
        db.close()
        self.redirect(f'/ho/order/{order_id}')


class HOMarkPackedHandler(BaseHandler):
    @tornado.web.authenticated
    def post(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        order_id = self.get_argument('order_id')
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        db = get_db()
        db.execute(
            "UPDATE stock_orders SET status='packed', packed_at=? WHERE id=? AND pack_date BETWEEN ? AND ?",
            (sast_now(), order_id, monday, sunday)
        )
        db.commit()
        db.close()
        self.redirect(f'/ho/week/{monday}')

class HOPackingListHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import tempfile
        except ImportError:
            self.write('<html><body style="font-family:Arial;padding:2rem">'
                      '<h2 style="color:#c0392b">python-docx nie geïnstalleer nie</h2>'
                      '<p>Hardloop: <code>pip install python-docx</code></p>'
                      '<a href="/ho/dashboard">Terug</a></body></html>')
            return

        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        db = get_db()
        SALON_ORDER = ['AV','BB','BV','EL','HE','PO','RS','TR','VG','WK']
        orders_raw = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.pack_date BETWEEN ? AND ? AND o.status IN ('submitted','packed')
            ORDER BY o.submitted_at DESC
        """, (monday, sunday)).fetchall()
        # Keep only the latest submitted/packed order per salon (skip older duplicates)
        salon_orders = {}
        for o in orders_raw:
            if o['salon_code'] not in salon_orders:
                salon_orders[o['salon_code']] = dict(o)

        salon_items = {}
        for code, order in salon_orders.items():
            rows = db.execute(
                "SELECT * FROM order_items WHERE order_id=? ORDER BY category, product_name",
                (order['id'],)
            ).fetchall()
            salon_items[code] = [dict(r) for r in rows]
        db.close()

        # ── Colour palette ───────────────────────────────────────────────────
        PURPLE  = RGBColor(0x3D, 0x1F, 0x5C)
        PURPLE2 = RGBColor(0x7B, 0x3F, 0xA0)
        LILAC   = RGBColor(0xC3, 0x9B, 0xD3)
        PALE    = RGBColor(0xED, 0xE0, 0xF8)
        SILVER  = RGBColor(0xF4, 0xF1, 0xF8)
        WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
        BORDER  = 'D5C8E8'
        CB      = '☐'

        # ── Helper functions ─────────────────────────────────────────────────

        def shd(cell, rgb):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            s  = OxmlElement('w:shd')
            s.set(qn('w:val'),   'clear')
            s.set(qn('w:color'), 'auto')
            s.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
            pr.append(s)

        def borders(cell, color=BORDER, sz='4',
                    sides=('top','bottom','left','right')):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            b  = OxmlElement('w:tcBorders')
            for side in sides:
                e = OxmlElement(f'w:{side}')
                e.set(qn('w:val'),   'single')
                e.set(qn('w:sz'),    sz)
                e.set(qn('w:space'), '0')
                e.set(qn('w:color'), color)
                b.append(e)
            pr.append(b)

        def no_borders(cell):
            borders(cell, color='FFFFFF', sz='0')

        def cell_para(cell, space_before=2, space_after=2):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            return p

        def fmt(qty):
            return str(int(qty)) if qty == int(qty) else str(qty)

        def lock_tbl_width(tbl, total_cm):
            """Force Word to respect absolute column widths by locking the table total width.
            Without this, Word redistributes column widths in multi-column sections."""
            tblPr = tbl._tbl.tblPr
            for old in tblPr.findall(qn('w:tblW')):
                tblPr.remove(old)
            tw = OxmlElement('w:tblW')
            tw.set(qn('w:w'), str(int(total_cm * 1440 / 2.54)))  # cm → twips
            tw.set(qn('w:type'), 'dxa')
            tblPr.append(tw)

        def force_col_widths(tbl, widths_cm):
            """Fix column widths in BOTH w:tblGrid and every cell's w:tcW.
            python-docx only writes cell-level widths; Word ignores them unless
            the tblGrid also matches — causing equal-width columns."""
            twips = [int(w * 1440 / 2.54) for w in widths_cm]
            # 1. Update w:tblGrid gridCol entries
            tblGrid = tbl._tbl.tblGrid
            grid_cols = tblGrid.findall(qn('w:gridCol'))
            for i, tw in enumerate(twips):
                if i < len(grid_cols):
                    grid_cols[i].set(qn('w:w'), str(tw))
                else:
                    gc = OxmlElement('w:gridCol')
                    gc.set(qn('w:w'), str(tw))
                    tblGrid.append(gc)
            # 2. Update every cell's w:tcW
            for row in tbl.rows:
                for ci, cell in enumerate(row.cells):
                    if ci >= len(widths_cm):
                        break
                    tcPr = cell._tc.get_or_add_tcPr()
                    for old in tcPr.findall(qn('w:tcW')):
                        tcPr.remove(old)
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(twips[ci]))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)
            # 3. Lock total table width
            lock_tbl_width(tbl, sum(widths_cm))

        def page_break(doc):
            p   = doc.add_paragraph()
            run = p.add_run()
            br  = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._r.append(br)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

        # ── Document setup — Landscape A4, 3-column ─────────────────────────
        doc  = Document()
        norm = doc.styles['Normal']
        norm.font.name = 'Calibri'
        norm.font.size = Pt(9)

        sec = doc.sections[0]
        sec.page_width    = Cm(29.7)   # Landscape: wide
        sec.page_height   = Cm(21.0)   # Landscape: short
        sec.left_margin   = Cm(0.8)
        sec.right_margin  = Cm(0.8)
        sec.top_margin    = Cm(0.8)
        sec.bottom_margin = Cm(0.8)

        # 3 equal columns, ~0.64cm gap between them
        # Per-column usable ≈ (29.7 - 1.6 - 2×0.64) / 3 ≈ 8.94cm
        sectPr  = sec._sectPr
        cols_el = OxmlElement('w:cols')
        cols_el.set(qn('w:num'),        '3')
        cols_el.set(qn('w:space'),      '360')   # 360 twips ≈ 0.64cm between cols
        cols_el.set(qn('w:equalWidth'), '1')
        sectPr.append(cols_el)

        # Item table column widths — 4 cols
        # Product(big) | Qty(medium) | Pkd | Del
        ITEM_WIDTHS = [5.0, 1.5, 1.0, 1.0]   # cm  → total 8.5 cm
        NCOLS       = 4

        TICK_BG = RGBColor(0xED, 0xE0, 0xF8)   # light lilac for checkbox cells

        def set_item_col_widths(tbl):
            force_col_widths(tbl, ITEM_WIDTHS)

        STATUS = {'submitted': 'Submitted', 'packed': 'Packed ✓', 'delivered': 'Delivered'}

        def col_break():
            """Insert a Word column-break so the next salon starts in the next column."""
            p   = doc.add_paragraph()
            run = p.add_run()
            br  = OxmlElement('w:br')
            br.set(qn('w:type'), 'column')
            run._r.append(br)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

        # ── Per-salon columns ─────────────────────────────────────────────────
        active_codes = [c for c in SALON_ORDER if c in salon_orders]

        for idx, code in enumerate(active_codes):
            order  = salon_orders[code]
            items  = salon_items.get(code, [])
            by_cat = {}
            for item in items:
                by_cat.setdefault(item['category'], []).append(item)
            # Count only real items (non-towel + towels with qty > 0)
            item_count = sum(
                1 for i in items
                if i['category'] != 'Handoeke' or float(i.get('quantity') or 0) > 0
            )

            # ── Compact salon header (2-cell table) ───────────────────────────
            hdr = doc.add_table(rows=1, cols=2)
            hdr.style   = 'Table Grid'
            hdr.autofit = False
            force_col_widths(hdr, [2.0, 6.5])

            # Left: large salon code
            c0 = hdr.rows[0].cells[0]
            shd(c0, PURPLE); borders(c0)
            p0 = cell_para(c0, 3, 3)
            r0 = p0.add_run(code)
            r0.bold = True; r0.font.size = Pt(20); r0.font.color.rgb = WHITE

            # Right: name + date + status
            c1 = hdr.rows[0].cells[1]
            shd(c1, PALE); borders(c1)
            p1 = cell_para(c1, 3, 3)
            nm = p1.add_run(order['salon_name'])
            nm.bold = True; nm.font.size = Pt(8); nm.font.color.rgb = PURPLE
            p1.add_run(f'\n{pack_date}  ·  {item_count} items').font.size = Pt(7)
            st_run = p1.add_run(f'\n{STATUS.get(order["status"], order["status"])}')
            st_run.bold = True; st_run.font.size = Pt(7.5); st_run.font.color.rgb = PURPLE2

            # Tiny spacer
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(1)
            sp.paragraph_format.space_after  = Pt(1)

            # ── Column-header row ──────────────────────────────────────────────
            col_hdr = doc.add_table(rows=1, cols=NCOLS)
            col_hdr.style   = 'Table Grid'
            col_hdr.autofit = False
            set_item_col_widths(col_hdr)
            for ci, (lbl, align, h_bg, h_border) in enumerate([
                ('Product',   WD_ALIGN_PARAGRAPH.LEFT,   PURPLE,  '3D1F5C'),
                ('Qty',       WD_ALIGN_PARAGRAPH.CENTER, PURPLE,  '3D1F5C'),
                ('Pkd',  WD_ALIGN_PARAGRAPH.CENTER, PURPLE2, '7B3FA0'),
                ('Del',  WD_ALIGN_PARAGRAPH.CENTER, PURPLE2, '7B3FA0'),
            ]):
                cell = col_hdr.rows[0].cells[ci]
                shd(cell, h_bg); borders(cell, color=h_border)
                p = cell_para(cell, 2, 1)
                p.alignment = align
                r = p.add_run(lbl)
                r.bold = True; r.font.size = Pt(7); r.font.color.rgb = WHITE
                if ci >= 2:
                    p.add_run(f' {CB}').font.size = Pt(9)

            # ── Category sections ──────────────────────────────────────────────
            for cat in ['Salon', 'Cleaning', 'Perms', 'Tints', 'Retail', 'Handoeke']:
                cat_items = by_cat.get(cat, [])
                # Filter out zero-qty towels (all towel types are stored, even empty ones)
                if cat == 'Handoeke':
                    cat_items = [i for i in cat_items if float(i.get('quantity') or 0) > 0]
                if not cat_items:
                    continue

                # Category header (merged full-width row)
                cat_tbl = doc.add_table(rows=1, cols=NCOLS)
                cat_tbl.style   = 'Table Grid'
                cat_tbl.autofit = False
                set_item_col_widths(cat_tbl)
                merged = cat_tbl.rows[0].cells[0].merge(cat_tbl.rows[0].cells[NCOLS-1])
                shd(merged, PURPLE2); borders(merged, color='7B3FA0', sz='4')
                p_cat = cell_para(merged, 2, 2)
                r_cat = p_cat.add_run(f' {cat.upper()}')
                r_cat.bold = True; r_cat.font.size = Pt(7.5); r_cat.font.color.rgb = WHITE

                # Item rows
                item_tbl = doc.add_table(rows=len(cat_items), cols=NCOLS)
                item_tbl.style   = 'Table Grid'
                item_tbl.autofit = False
                set_item_col_widths(item_tbl)

                for ri, item in enumerate(cat_items):
                    bg  = WHITE if ri % 2 == 0 else SILVER
                    row = item_tbl.rows[ri]

                    # Product name
                    c_nm = row.cells[0]
                    shd(c_nm, bg); borders(c_nm)
                    p_nm = cell_para(c_nm, 2, 2)
                    nm_r = p_nm.add_run(item['product_name'])
                    nm_r.font.size = Pt(8)
                    if item.get('note'):
                        nt = p_nm.add_run(f'\n{item["note"]}')
                        nt.font.size = Pt(6.5); nt.font.color.rgb = RGBColor(0x7A,0x6B,0x8A)

                    # Hoeveelheid
                    c_qt = row.cells[1]
                    shd(c_qt, bg); borders(c_qt)
                    p_qt = cell_para(c_qt, 2, 2)
                    p_qt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    qt_r = p_qt.add_run(f'×{fmt(item["quantity"])}')
                    qt_r.bold = True; qt_r.font.size = Pt(9); qt_r.font.color.rgb = PURPLE2

                    # Gepak ☐ — lilac bg so the tick box is clearly visible
                    c_gp = row.cells[2]
                    shd(c_gp, TICK_BG); borders(c_gp, color='C39BD3', sz='4')
                    p_gp = cell_para(c_gp, 2, 2)
                    p_gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_gp.add_run(CB).font.size = Pt(12)

                    # Afgelewer ☐ — lilac bg
                    c_af = row.cells[3]
                    shd(c_af, TICK_BG); borders(c_af, color='C39BD3', sz='4')
                    p_af = cell_para(c_af, 2, 2)
                    p_af.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_af.add_run(CB).font.size = Pt(12)

            # ── Compact signature line ─────────────────────────────────────────
            sp2 = doc.add_paragraph()
            sp2.paragraph_format.space_before = Pt(2)
            sp2.paragraph_format.space_after  = Pt(2)
            sig_tbl = doc.add_table(rows=1, cols=2)
            sig_tbl.style   = 'Table Grid'
            sig_tbl.autofit = False
            force_col_widths(sig_tbl, [4.25, 4.25])
            for ci, lbl in enumerate(['Received by:', 'Date:']):
                cell = sig_tbl.rows[0].cells[ci]
                shd(cell, PALE); borders(cell)
                p = cell_para(cell, 3, 3)
                r1 = p.add_run(lbl + '  ')
                r1.bold = True; r1.font.size = Pt(7.5); r1.font.color.rgb = PURPLE
                p.add_run('_' * 16).font.size = Pt(7.5)

            # ── Column break (not after the very last salon) ───────────────────
            if idx < len(active_codes) - 1:
                col_break()

        # ── Save & send ──────────────────────────────────────────────────────
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, 'rb') as f:
            data = f.read()
        os.unlink(tmp.name)

        fname = f"SVS_Pakkingslys_{pack_date}.docx"
        self.set_header('Content-Type',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.set_header('Content-Disposition', f'attachment; filename="{fname}"')
        self.write(data)


class HOPrintViewHandler(BaseHandler):
    """Browser-drukbare gesamentlike lys of uitstaande lys vir 'n week."""
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        view = self.get_argument('type', 'all')   # 'all' or 'outstanding'
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        db = get_db()
        orders = db.execute("""
            SELECT o.*, s.code as salon_code, s.name as salon_name
            FROM stock_orders o JOIN salons s ON o.salon_id = s.id
            WHERE o.pack_date BETWEEN ? AND ? AND o.status != 'draft'
            ORDER BY s.code
        """, (monday, sunday)).fetchall()

        items_raw = db.execute("""
            SELECT oi.product_name, oi.category, oi.quantity,
                   COALESCE(oi.packed, 0) as packed,
                   CASE WHEN COALESCE(oi.packed,0)=1
                        THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                        ELSE 0 END as item_packed_qty,
                   s.code as salon_code
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            JOIN salons s ON so.salon_id = s.id
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
            ORDER BY oi.category, oi.product_name
        """, (monday, sunday)).fetchall()
        db.close()

        from collections import OrderedDict
        agg = OrderedDict()
        for row in items_raw:
            key = (row['category'], row['product_name'])
            if key not in agg:
                agg[key] = {'category': row['category'], 'product_name': row['product_name'],
                            'total_qty': 0, 'packed_qty': 0,
                            'packed_salons': 0, 'total_salons': 0, 'salons': []}
            qty = row['quantity']
            pq  = row['item_packed_qty']
            agg[key]['total_qty']  += qty
            agg[key]['packed_qty'] += pq
            existing = next((s for s in agg[key]['salons'] if s['code'] == row['salon_code']), None)
            if existing:
                existing['qty']        += qty
                existing['packed_qty'] += pq
                if not row['packed']:
                    existing['packed'] = False
            else:
                agg[key]['salons'].append({
                    'code':       row['salon_code'],
                    'qty':        qty,
                    'packed_qty': pq,
                    'packed':     bool(row['packed'])
                })
                agg[key]['total_salons'] += 1
                if row['packed']:
                    agg[key]['packed_salons'] += 1

        all_items   = list(agg.values())
        outstanding = [i for i in all_items
                       if i['packed_qty'] < i['total_qty']
                       and i['category'] != 'Handoeke']
        show_items  = outstanding if view == 'outstanding' else all_items
        title       = 'Outstanding Items' if view == 'outstanding' else 'Combined Packing List'

        self.render('ho/print_view.html', user=u, pack_date=monday,
                    week_friday=week_friday(monday),
                    orders=[dict(o) for o in orders], items=show_items,
                    view=view, title=title)


class HOOutstandingDocHandler(BaseHandler):
    """Genereer die uitstaande items as Word-dokument gegroepeer per kategorie."""
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            self.write('<html><body style="font-family:Arial;padding:2rem">'
                       '<h2 style="color:#c0392b">python-docx nie geïnstalleer nie</h2>'
                       '<p>Hardloop: <code>pip install python-docx</code></p>'
                       '<a href="/ho/dashboard">Terug</a></body></html>')
            return

        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        friday = week_friday(monday)
        db = get_db()
        items_raw = db.execute("""
            SELECT oi.product_name, oi.category, oi.quantity,
                   CASE WHEN COALESCE(oi.packed,0)=1
                        THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                        ELSE 0 END as item_packed_qty,
                   s.code as salon_code
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            JOIN salons s ON so.salon_id = s.id
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category != 'Handoeke'
            ORDER BY oi.category, oi.product_name, s.code
        """, (monday, sunday)).fetchall()
        db.close()

        from collections import OrderedDict
        agg = OrderedDict()
        for row in items_raw:
            key = (row['category'], row['product_name'])
            if key not in agg:
                agg[key] = {
                    'product_name': row['product_name'],
                    'category': row['category'],
                    'total_qty': 0, 'packed_qty': 0, 'salons': []
                }
            qty = row['quantity']
            pq  = row['item_packed_qty']
            agg[key]['total_qty']  += qty
            agg[key]['packed_qty'] += pq
            existing = next((s for s in agg[key]['salons'] if s['code'] == row['salon_code']), None)
            if existing:
                existing['qty']        += qty
                existing['packed_qty'] += pq
            else:
                agg[key]['salons'].append({'code': row['salon_code'], 'qty': qty, 'packed_qty': pq})

        outstanding_items = [v for v in agg.values() if v['packed_qty'] < v['total_qty']]

        # ── Colour palette ────────────────────────────────────────────────────
        RED   = RGBColor(0xC0, 0x39, 0x2B)  # #C0392B
        LRED  = RGBColor(0xFD, 0xEB, 0xD0)  # #FDEBD0 warm light
        PALE  = RGBColor(0xFF, 0xFB, 0xF0)  # #FFFBF0
        DARK  = RGBColor(0x2C, 0x1A, 0x3E)  # #2C1A3E
        GREEN = RGBColor(0x1E, 0x84, 0x49)  # #1E8449
        WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        BORDER = 'F5B7B1'

        def shd(cell, rgb):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            s  = OxmlElement('w:shd')
            s.set(qn('w:val'),   'clear')
            s.set(qn('w:color'), 'auto')
            s.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
            pr.append(s)

        def borders(cell, color=BORDER, sz='4', sides=('top','bottom','left','right')):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            b  = OxmlElement('w:tcBorders')
            for side in sides:
                e = OxmlElement(f'w:{side}')
                e.set(qn('w:val'),   'single')
                e.set(qn('w:sz'),    sz)
                e.set(qn('w:space'), '0')
                e.set(qn('w:color'), color)
                b.append(e)
            pr.append(b)

        def no_borders(cell):
            borders(cell, color='FFFFFF', sz='0')

        def cell_para(cell, space_before=2, space_after=2):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            return p

        def fmt(qty):
            v = float(qty) if qty is not None else 0.0
            return str(int(v)) if v == int(v) else str(v)

        def lock_tbl_width(tbl, total_cm):
            tblPr = tbl._tbl.tblPr
            for old in tblPr.findall(qn('w:tblW')):
                tblPr.remove(old)
            tw = OxmlElement('w:tblW')
            tw.set(qn('w:w'), str(int(total_cm * 1440 / 2.54)))
            tw.set(qn('w:type'), 'dxa')
            tblPr.append(tw)

        def force_col_widths(tbl, widths_cm):
            twips = [int(w * 1440 / 2.54) for w in widths_cm]
            tblGrid = tbl._tbl.tblGrid
            grid_cols = tblGrid.findall(qn('w:gridCol'))
            for i, tw in enumerate(twips):
                if i < len(grid_cols):
                    grid_cols[i].set(qn('w:w'), str(tw))
                else:
                    gc = OxmlElement('w:gridCol')
                    gc.set(qn('w:w'), str(tw))
                    tblGrid.append(gc)
            for row in tbl.rows:
                for ci, cell in enumerate(row.cells):
                    if ci >= len(widths_cm): break
                    tcPr = cell._tc.get_or_add_tcPr()
                    for old in tcPr.findall(qn('w:tcW')):
                        tcPr.remove(old)
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(twips[ci]))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)
            lock_tbl_width(tbl, sum(widths_cm))

        # ── Document setup ────────────────────────────────────────────────────
        doc  = Document()
        norm = doc.styles['Normal']
        norm.font.name = 'Calibri'
        norm.font.size = Pt(10)

        sec = doc.sections[0]
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)

        # Usable: 17.4 cm total
        W_CB   = 0.7
        W_PROD = 6.0
        W_OUT  = 1.8
        W_SAL  = 8.9

        # ── Page header ───────────────────────────────────────────────────────
        hdr = doc.add_table(rows=1, cols=2)
        hdr.style = 'Table Grid'; hdr.alignment = WD_TABLE_ALIGNMENT.LEFT; hdr.autofit = False
        force_col_widths(hdr, [10.5, 6.9])

        c0 = hdr.rows[0].cells[0]; shd(c0, RED); borders(c0)
        p0 = cell_para(c0, 5, 2)
        r0 = p0.add_run('SVS Stock'); r0.bold = True; r0.font.size = Pt(18); r0.font.color.rgb = WHITE
        r0b = p0.add_run('  ·  Uitstaande Items'); r0b.font.size = Pt(10); r0b.font.color.rgb = LRED
        p0s = c0.add_paragraph(); p0s.paragraph_format.space_before = Pt(1); p0s.paragraph_format.space_after = Pt(5)
        r0c = p0s.add_run('Silver Violet Studios – items nog nie gepak nie')
        r0c.font.size = Pt(8); r0c.font.color.rgb = LRED

        c1 = hdr.rows[0].cells[1]; shd(c1, LRED); borders(c1)
        p1 = cell_para(c1, 5, 2); p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(f'Week: {monday} → {friday}')
        r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = RED
        p1s = c1.add_paragraph(); p1s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1s.paragraph_format.space_before = Pt(2); p1s.paragraph_format.space_after = Pt(5)
        r1b = p1s.add_run(f'Gegenereer: {sast_now()[:16]}')
        r1b.font.size = Pt(7.5); r1b.font.color.rgb = RED

        sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(10); sp.paragraph_format.space_after = Pt(0)

        # ── Per-category sections ─────────────────────────────────────────────
        for cat in ['Salon', 'Cleaning', 'Perms', 'Tints', 'Retail']:
            cat_items = [i for i in outstanding_items if i['category'] == cat]
            if not cat_items:
                continue

            # Spacer between categories
            sp2 = doc.add_paragraph(); sp2.paragraph_format.space_before = Pt(8); sp2.paragraph_format.space_after = Pt(0)

            # Category heading
            ch = doc.add_table(rows=1, cols=1)
            ch.style = 'Table Grid'; ch.alignment = WD_TABLE_ALIGNMENT.LEFT; ch.autofit = False
            force_col_widths(ch, [17.4])
            cc = ch.rows[0].cells[0]; shd(cc, RED); borders(cc, color='C0392B')
            cp = cell_para(cc, 3, 3)
            cr = cp.add_run(f'  {cat.upper()}')
            cr.bold = True; cr.font.size = Pt(10); cr.font.color.rgb = WHITE
            cnt = len(cat_items)
            cr2 = cp.add_run(f'    ({cnt} item{"s" if cnt != 1 else ""})')
            cr2.font.size = Pt(8); cr2.font.color.rgb = LRED

            # Table
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; tbl.autofit = False
            force_col_widths(tbl, [W_CB, W_PROD, W_OUT, W_SAL])

            # Header row
            hrow = tbl.rows[0]
            for ci, (htxt, halign) in enumerate([
                ('☐', WD_ALIGN_PARAGRAPH.CENTER),
                ('Produk / Product', WD_ALIGN_PARAGRAPH.LEFT),
                ('Uitstaande', WD_ALIGN_PARAGRAPH.CENTER),
                ('Per Salon', WD_ALIGN_PARAGRAPH.LEFT),
            ]):
                hc = hrow.cells[ci]; shd(hc, LRED); borders(hc, color=BORDER)
                hp = cell_para(hc, 2, 2); hp.alignment = halign
                hr = hp.add_run(htxt); hr.bold = True; hr.font.size = Pt(8); hr.font.color.rgb = RED

            # Data rows
            for ri, item in enumerate(cat_items):
                drow = tbl.add_row()
                bg = LRED if ri % 2 == 0 else WHITE
                out_qty = item['total_qty'] - item['packed_qty']

                c_cb = drow.cells[0]; shd(c_cb, bg); borders(c_cb, color=BORDER)
                pcb = cell_para(c_cb, 3, 3); pcb.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pcb.add_run('☐').font.size = Pt(11)

                c_pr = drow.cells[1]; shd(c_pr, bg); borders(c_pr, color=BORDER)
                ppr = cell_para(c_pr, 3, 3)
                rpr = ppr.add_run(item['product_name'])
                rpr.bold = True; rpr.font.size = Pt(9); rpr.font.color.rgb = DARK

                c_ot = drow.cells[2]; shd(c_ot, bg); borders(c_ot, color=BORDER)
                pot = cell_para(c_ot, 3, 3); pot.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rot = pot.add_run(fmt(out_qty))
                rot.bold = True; rot.font.size = Pt(11); rot.font.color.rgb = RED

                c_sl = drow.cells[3]; shd(c_sl, bg); borders(c_sl, color=BORDER)
                psl = cell_para(c_sl, 3, 3)
                salon_parts = []
                for s in item['salons']:
                    s_out = s['qty'] - s['packed_qty']
                    if s_out > 0:
                        salon_parts.append(f"{s['code']}:{fmt(s_out)}")
                rsl = psl.add_run('  '.join(salon_parts))
                rsl.font.size = Pt(8.5); rsl.font.color.rgb = DARK

        if not outstanding_items:
            done_p = doc.add_paragraph()
            done_p.paragraph_format.space_before = Pt(20)
            dr = done_p.add_run('✅  Alles gepak — geen uitstaande items nie!')
            dr.bold = True; dr.font.size = Pt(12); dr.font.color.rgb = GREEN

        # ── Footer ────────────────────────────────────────────────────────────
        fsp = doc.add_paragraph(); fsp.paragraph_format.space_before = Pt(14); fsp.paragraph_format.space_after = Pt(0)
        f_tbl = doc.add_table(rows=1, cols=2)
        f_tbl.style = 'Table Grid'; f_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; f_tbl.autofit = False
        force_col_widths(f_tbl, [11.4, 6.0])
        fc0 = f_tbl.rows[0].cells[0]; fc1 = f_tbl.rows[0].cells[1]
        shd(fc0, LRED); no_borders(fc0); shd(fc1, LRED); no_borders(fc1)
        fp0 = cell_para(fc0, 3, 3)
        fp0.add_run('Silver Violet Studios – SVS Stock System').font.size = Pt(7.5)
        fp1 = cell_para(fc1, 3, 3); fp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rfp = fp1.add_run(f'Week: {monday}'); rfp.font.size = Pt(7.5); rfp.font.color.rgb = RED

        # ── Serve as download ─────────────────────────────────────────────────
        import tempfile as _tmpfile, os as _os
        with _tmpfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        doc.save(tmp_path)
        with open(tmp_path, 'rb') as f:
            data = f.read()
        _os.unlink(tmp_path)

        fname = f'SVS_Uitstaande_{monday}.docx'
        self.set_header('Content-Type',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.set_header('Content-Disposition', f'attachment; filename="{fname}"')
        self.write(data)


class HOOutstandingHandler(BaseHandler):
    """Interaktiewe uitstaande-bestel-bladsy — Sarie merk wat sy by die verskaffer bestel het."""
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        db = get_db()
        items_raw = db.execute("""
            SELECT oi.product_name, oi.category,
                   SUM(oi.quantity) as total_qty,
                   SUM(CASE WHEN COALESCE(oi.packed,0)=1 THEN oi.quantity ELSE 0 END) as packed_qty
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category != 'Handoeke'
            GROUP BY oi.product_name, oi.category
            HAVING SUM(CASE WHEN COALESCE(oi.packed,0)=0 THEN 1 ELSE 0 END) > 0
            ORDER BY oi.category, oi.product_name
        """, (monday, sunday)).fetchall()

        # Laai bestaande supplier_orders vir hierdie week (gebruik Maandag as sleutel)
        existing = db.execute(
            "SELECT * FROM supplier_orders WHERE pack_date=?", (monday,)
        ).fetchall()
        db.close()

        sup = {r['product_name']: dict(r) for r in existing}
        outstanding = []
        for row in items_raw:
            qty_out = row['total_qty'] - row['packed_qty']
            rec = sup.get(row['product_name'], {})
            outstanding.append({
                'product_name': row['product_name'],
                'category':     row['category'],
                'total_qty':    row['total_qty'],
                'packed_qty':   row['packed_qty'],
                'qty_needed':   qty_out,
                'ordered':      rec.get('ordered', 0),
                'qty_ordered':  rec.get('qty_ordered', qty_out),
                'notes':        rec.get('notes', ''),
            })

        self.render('ho/outstanding.html', user=u, pack_date=monday,
                    outstanding=outstanding)

    @tornado.web.authenticated
    def post(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        db = get_db()
        items_raw = db.execute("""
            SELECT oi.product_name, oi.category,
                   SUM(oi.quantity) as total_qty,
                   SUM(CASE WHEN COALESCE(oi.packed,0)=1 THEN oi.quantity ELSE 0 END) as packed_qty
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category != 'Handoeke'
            GROUP BY oi.product_name, oi.category
            HAVING SUM(CASE WHEN COALESCE(oi.packed,0)=0 THEN 1 ELSE 0 END) > 0
            ORDER BY oi.category, oi.product_name
        """, (monday, sunday)).fetchall()

        # Wis ou rekords vir hierdie week uit en skryf nuut (gebruik Maandag as sleutel)
        db.execute("DELETE FROM supplier_orders WHERE pack_date=?", (monday,))
        for i, row in enumerate(items_raw):
            ordered = 1 if self.get_argument(f'ord_{i}', '') == 'on' else 0
            try:
                qty_ord = float(self.get_argument(f'qty_{i}', str(row['total_qty'] - row['packed_qty'])))
            except Exception:
                qty_ord = row['total_qty'] - row['packed_qty']
            notes = self.get_argument(f'notes_{i}', '').strip()
            ordered_at = sast_now() if ordered else None
            db.execute("""
                INSERT INTO supplier_orders
                  (pack_date, product_name, category, qty_needed, qty_ordered, ordered, ordered_at, notes)
                VALUES (?,?,?,?,?,?,?,?)
            """, (monday, row['product_name'], row['category'],
                  row['total_qty'] - row['packed_qty'],
                  qty_ord, ordered, ordered_at, notes))
        db.commit()
        db.close()
        self.redirect(f'/ho/week/{monday}')


# ─── Towel handlers ──────────────────────────────────────────────────────────

class HOTowelPrintHandler(BaseHandler):
    """Browser-drukbare handoek-opsommingslys vir 'n week."""
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        friday = week_friday(monday)
        db = get_db()
        SALON_ORDER = ['AV','BB','BV','EL','HE','PO','RS','TR','VG','WK']
        items_raw = db.execute("""
            SELECT oi.product_name, oi.quantity, s.code as salon_code
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            JOIN salons s ON so.salon_id = s.id
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category = 'Handoeke' AND oi.quantity > 0
            ORDER BY oi.product_name, s.code
        """, (monday, sunday)).fetchall()
        db.close()

        from collections import OrderedDict
        matrix = OrderedDict()
        salon_set = set()
        for row in items_raw:
            tt = row['product_name']
            if tt not in matrix:
                matrix[tt] = {}
            matrix[tt][row['salon_code']] = float(row['quantity'])
            salon_set.add(row['salon_code'])

        towel_types  = list(matrix.keys())
        towel_salons = [c for c in SALON_ORDER if c in salon_set]
        self.render('ho/towel_print.html', user=u,
                    pack_date=monday, week_friday=friday,
                    towel_types=towel_types, towel_salons=towel_salons,
                    towel_matrix=matrix)


class HOTowelWordHandler(BaseHandler):
    """Genereer handoek-opsomming as Word-dokument (matriksformaat)."""
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import tempfile
        except ImportError:
            self.write('<html><body style="font-family:Arial;padding:2rem">'
                       '<h2 style="color:#c0392b">python-docx nie geïnstalleer nie</h2>'
                       '<p>Hardloop: <code>pip install python-docx</code></p>'
                       '<a href="/ho/dashboard">Terug</a></body></html>')
            return

        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        friday = week_friday(monday)
        db = get_db()
        SALON_ORDER = ['AV','BB','BV','EL','HE','PO','RS','TR','VG','WK']
        items_raw = db.execute("""
            SELECT oi.product_name, oi.quantity, s.code as salon_code
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            JOIN salons s ON so.salon_id = s.id
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category = 'Handoeke' AND oi.quantity > 0
            ORDER BY oi.product_name, s.code
        """, (monday, sunday)).fetchall()
        db.close()

        from collections import OrderedDict
        matrix = OrderedDict()
        salon_set = set()
        for row in items_raw:
            tt = row['product_name']
            if tt not in matrix:
                matrix[tt] = {}
            matrix[tt][row['salon_code']] = float(row['quantity'])
            salon_set.add(row['salon_code'])

        towel_types  = list(matrix.keys())
        towel_salons = [c for c in SALON_ORDER if c in salon_set]

        def fmt(v):
            v = float(v) if v is not None else 0.0
            return str(int(v)) if v == int(v) else str(v)

        # ── Colour palette ────────────────────────────────────────────────────
        PURPLE  = RGBColor(0x3D, 0x1F, 0x5C)
        PURPLE2 = RGBColor(0x7B, 0x3F, 0xA0)
        PALE    = RGBColor(0xED, 0xE0, 0xF8)
        WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
        SILVER  = RGBColor(0xF4, 0xF1, 0xF8)
        BORDER  = 'D5C8E8'

        def shd(cell, rgb):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            s = OxmlElement('w:shd')
            s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
            s.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
            pr.append(s)

        def borders(cell, color=BORDER, sz='4', sides=('top','bottom','left','right')):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            b = OxmlElement('w:tcBorders')
            for side in sides:
                e = OxmlElement(f'w:{side}')
                e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
                e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
                b.append(e)
            pr.append(b)

        def set_cell_width(cell, cm):
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(cm * 1440 / 2.54)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        def cp(cell, sb=2, sa=2):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after  = Pt(sa)
            return p

        doc  = Document()
        norm = doc.styles['Normal']
        norm.font.name = 'Calibri'; norm.font.size = Pt(10)

        sec = doc.sections[0]
        # Landscape for wide matrix
        sec.page_width    = Cm(29.7)
        sec.page_height   = Cm(21.0)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)

        usable = 29.7 - 3.0  # 26.7 cm usable
        W_TYPE  = 3.8   # Towel type column
        W_TOTAL = 1.5   # Total column
        n_sal   = len(towel_salons)
        W_SAL   = round((usable - W_TYPE - W_TOTAL) / n_sal, 2) if n_sal else 2.0

        # ── Page header ───────────────────────────────────────────────────────
        hdr = doc.add_table(rows=1, cols=2)
        hdr.style = 'Table Grid'; hdr.autofit = False
        hdr.columns[0].width = Cm(16); hdr.columns[1].width = Cm(10.7)

        c0 = hdr.rows[0].cells[0]; shd(c0, PURPLE); borders(c0)
        p0 = cp(c0, 5, 2)
        r0 = p0.add_run('SVS Stock'); r0.bold = True; r0.font.size = Pt(16); r0.font.color.rgb = WHITE
        r0b = p0.add_run('  ·  Handoeke / Towels'); r0b.font.size = Pt(10); r0b.font.color.rgb = PALE
        p0b = c0.add_paragraph(); p0b.paragraph_format.space_before = Pt(0); p0b.paragraph_format.space_after = Pt(5)
        p0b.add_run(f'Silver Violet Studios  –  {monday} tot {friday}').font.size = Pt(8)
        p0b.runs[-1].font.color.rgb = PALE

        c1 = hdr.rows[0].cells[1]; shd(c1, PALE); borders(c1)
        p1 = cp(c1, 5, 2)
        r1 = p1.add_run(f'Pak-week: {monday} – {friday}')
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = PURPLE
        p1.add_run(f'\n{len(towel_types)} handoek-tipes  ·  {len(towel_salons)} salonne')

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(3)
        sp.paragraph_format.space_after  = Pt(3)

        if not towel_types:
            p = doc.add_paragraph()
            p.add_run('Geen handoek-bestellings vir hierdie week nie.').italic = True
        else:
            # ── Matrix table ──────────────────────────────────────────────────
            ncols = 1 + 1 + n_sal   # Type + Total + salons
            tbl = doc.add_table(rows=1 + len(towel_types) + 1, cols=ncols)
            tbl.style = 'Table Grid'; tbl.autofit = False

            # Column widths
            set_cell_width(tbl.rows[0].cells[0], W_TYPE)
            set_cell_width(tbl.rows[0].cells[1], W_TOTAL)
            for ci in range(n_sal):
                set_cell_width(tbl.rows[0].cells[2 + ci], W_SAL)

            # Header row
            hrow = tbl.rows[0]
            shd(hrow.cells[0], PURPLE); borders(hrow.cells[0])
            ph = cp(hrow.cells[0], 3, 2)
            rh = ph.add_run('Handoek Tipe')
            rh.bold = True; rh.font.size = Pt(9); rh.font.color.rgb = WHITE

            shd(hrow.cells[1], PURPLE2); borders(hrow.cells[1])
            ph2 = cp(hrow.cells[1], 3, 2); ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rh2 = ph2.add_run('Totaal')
            rh2.bold = True; rh2.font.size = Pt(9); rh2.font.color.rgb = WHITE

            for ci, code in enumerate(towel_salons):
                cell = hrow.cells[2 + ci]
                shd(cell, PURPLE2); borders(cell)
                ph = cp(cell, 3, 2); ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rh = ph.add_run(code)
                rh.bold = True; rh.font.size = Pt(9); rh.font.color.rgb = WHITE

            # Data rows
            for ri, tt in enumerate(towel_types):
                row  = tbl.rows[1 + ri]
                bg   = WHITE if ri % 2 == 0 else SILVER
                row_data = matrix.get(tt, {})
                total_q  = sum(row_data.values())

                # Type name
                shd(row.cells[0], bg); borders(row.cells[0])
                p_name = cp(row.cells[0], 3, 2)
                r_name = p_name.add_run(tt)
                r_name.bold = True; r_name.font.size = Pt(10); r_name.font.color.rgb = PURPLE

                # Total
                shd(row.cells[1], bg); borders(row.cells[1])
                p_tot = cp(row.cells[1], 3, 2); p_tot.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_tot = p_tot.add_run(fmt(total_q))
                r_tot.bold = True; r_tot.font.size = Pt(11); r_tot.font.color.rgb = PURPLE2

                # Per salon
                for ci, code in enumerate(towel_salons):
                    cell = row.cells[2 + ci]
                    shd(cell, bg); borders(cell)
                    qty = row_data.get(code, 0)
                    p_sal = cp(cell, 3, 2); p_sal.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if qty > 0:
                        r_sal = p_sal.add_run(fmt(qty))
                        r_sal.font.size = Pt(10); r_sal.font.color.rgb = PURPLE
                    else:
                        r_sal = p_sal.add_run('–')
                        r_sal.font.size = Pt(9); r_sal.font.color.rgb = RGBColor(0xC0,0xC0,0xC0)

            # Totals row
            trow = tbl.rows[-1]
            shd(trow.cells[0], PALE); borders(trow.cells[0])
            p_tl = cp(trow.cells[0], 3, 3)
            r_tl = p_tl.add_run('TOTAAL')
            r_tl.bold = True; r_tl.font.size = Pt(9); r_tl.font.color.rgb = PURPLE

            grand_total = sum(sum(d.values()) for d in matrix.values())
            shd(trow.cells[1], PALE); borders(trow.cells[1])
            p_gt = cp(trow.cells[1], 3, 3); p_gt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_gt = p_gt.add_run(fmt(grand_total))
            r_gt.bold = True; r_gt.font.size = Pt(11); r_gt.font.color.rgb = PURPLE

            for ci, code in enumerate(towel_salons):
                cell = trow.cells[2 + ci]
                shd(cell, PALE); borders(cell)
                sal_total = sum(matrix.get(tt, {}).get(code, 0) for tt in towel_types)
                p_st = cp(cell, 3, 3); p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_st = p_st.add_run(fmt(sal_total))
                r_st.bold = True; r_st.font.size = Pt(10); r_st.font.color.rgb = PURPLE

        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(tmp.name); tmp.close()
        with open(tmp.name, 'rb') as f:
            data = f.read()
        os.unlink(tmp.name)

        fname = f"SVS_Handoeke_{monday}.docx"
        self.set_header('Content-Type',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.set_header('Content-Disposition', f'attachment; filename="{fname}"')
        self.write(data)


class HOOrderListHandler(BaseHandler):
    """Bestel-lys vir uitstaande items gegroepeer per verskaffer."""
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        db = get_db()
        items_raw = db.execute("""
            SELECT oi.product_name, oi.category,
                   SUM(oi.quantity)   as total_qty,
                   SUM(CASE WHEN COALESCE(oi.packed,0)=1
                            THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                            ELSE 0 END) as packed_qty,
                   COALESCE(p.supplier, '') as supplier
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            LEFT JOIN products p
                   ON LOWER(TRIM(p.name)) = LOWER(TRIM(oi.product_name))
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category != 'Handoeke'
            GROUP BY oi.product_name, oi.category, p.supplier
            HAVING SUM(oi.quantity) > SUM(
                CASE WHEN COALESCE(oi.packed,0)=1
                     THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                     ELSE 0 END)
            ORDER BY COALESCE(p.supplier,''), oi.category, oi.product_name
        """, (monday, sunday)).fetchall()
        # Read any saved supplier-order adjustments for this week
        sup_saved = {r['product_name']: dict(r) for r in
                     db.execute("SELECT * FROM supplier_orders WHERE pack_date=?",
                                (monday,)).fetchall()}

        # Pre-fetch supplier info for extra products (those in sup_saved but not outstanding)
        outstanding_names = {row['product_name'] for row in items_raw}
        extra_names = [pname for pname in sup_saved
                       if pname not in outstanding_names
                       and float(sup_saved[pname].get('qty_ordered') or 0) > 0]
        extra_prod_info = {}
        for pname in extra_names:
            row2 = db.execute(
                "SELECT COALESCE(supplier,'') as supplier, category FROM products "
                "WHERE LOWER(TRIM(name))=LOWER(TRIM(?))", (pname,)
            ).fetchone()
            if row2:
                extra_prod_info[pname] = {
                    'supplier': row2['supplier'] or ('HHB' if row2['category'] == 'Tints' else 'No Supplier'),
                    'category': row2['category'],
                }
            else:
                cat_saved = sup_saved[pname].get('category', 'Salon')
                # Koppel Tints-kategorie outomaties aan HHB
                sup_default = 'HHB' if cat_saved == 'Tints' else 'No Supplier'
                extra_prod_info[pname] = {
                    'supplier': sup_default,
                    'category': cat_saved,
                }
        db.close()

        from collections import OrderedDict
        suppliers = OrderedDict()
        for row in items_raw:
            sup = row['supplier'] or ('HHB' if row['category'] == 'Tints' else 'No Supplier')
            qty_out = row['total_qty'] - row['packed_qty']
            rec = sup_saved.get(row['product_name'], {})
            # Skip items where HO explicitly set qty to 0 (✕ gedruk en gestoor)
            if rec and float(rec.get('qty_ordered') or 0) == 0:
                continue
            if sup not in suppliers:
                suppliers[sup] = []
            suppliers[sup].append({
                'product_name': row['product_name'],
                'category':     row['category'],
                'qty_needed':   qty_out,
                'qty_ordered':  rec.get('qty_ordered', qty_out) if rec else qty_out,
                'notes':        rec.get('notes', '') if rec else '',
            })
        # Add extra manually-added products from supplier_orders
        for pname in extra_names:
            rec  = sup_saved[pname]
            info = extra_prod_info.get(pname, {'supplier': 'No Supplier', 'category': 'Salon'})
            sup_key = info['supplier']
            if sup_key not in suppliers:
                suppliers[sup_key] = []
            suppliers[sup_key].append({
                'product_name': pname,
                'category':     info['category'],
                'qty_needed':   0,
                'qty_ordered':  float(rec.get('qty_ordered') or 0),
                'notes':        rec.get('notes', '') or '',
            })

        self.render('ho/order_list.html', user=u, pack_date=monday,
                    week_friday=week_friday(monday),
                    suppliers=suppliers)


class HOOrderListDocHandler(BaseHandler):
    """Genereer die Bestelys as Word-dokument gegroepeer per verskaffer."""
    @tornado.web.authenticated
    def get(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            self.write('<html><body style="font-family:Arial;padding:2rem">'
                       '<h2 style="color:#c0392b">python-docx nie geïnstalleer nie</h2>'
                       '<p>Hardloop: <code>pip install python-docx</code></p>'
                       '<a href="/ho/dashboard">Terug</a></body></html>')
            return

        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        friday = week_friday(monday)
        db = get_db()
        items_raw = db.execute("""
            SELECT oi.product_name, oi.category,
                   SUM(oi.quantity)   as total_qty,
                   SUM(CASE WHEN COALESCE(oi.packed,0)=1
                            THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                            ELSE 0 END) as packed_qty,
                   COALESCE(p.supplier, '') as supplier
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            LEFT JOIN products p
                   ON LOWER(TRIM(p.name)) = LOWER(TRIM(oi.product_name))
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category != 'Handoeke'
            GROUP BY oi.product_name, oi.category, p.supplier
            HAVING SUM(oi.quantity) > SUM(
                CASE WHEN COALESCE(oi.packed,0)=1
                     THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                     ELSE 0 END)
            ORDER BY COALESCE(p.supplier,''), oi.category, oi.product_name
        """, (monday, sunday)).fetchall()
        sup_saved = {r['product_name']: dict(r) for r in
                     db.execute("SELECT * FROM supplier_orders WHERE pack_date=?",
                                (monday,)).fetchall()}
        outstanding_names = {row['product_name'] for row in items_raw}
        extra_names = [pname for pname in sup_saved
                       if pname not in outstanding_names
                       and float(sup_saved[pname].get('qty_ordered') or 0) > 0]
        extra_prod_info = {}
        for pname in extra_names:
            row2 = db.execute(
                "SELECT COALESCE(supplier,'') as supplier, category FROM products "
                "WHERE LOWER(TRIM(name))=LOWER(TRIM(?))", (pname,)
            ).fetchone()
            if row2:
                extra_prod_info[pname] = {
                    'supplier': row2['supplier'] or 'No Supplier',
                    'category': row2['category'],
                }
            else:
                cat_saved = sup_saved[pname].get('category', 'Salon')
                extra_prod_info[pname] = {
                    'supplier': 'HHB' if cat_saved == 'Tints' else 'No Supplier',
                    'category': cat_saved,
                }
        db.close()

        from collections import OrderedDict
        suppliers = OrderedDict()
        for row in items_raw:
            sup = row['supplier'] or 'No Supplier'
            qty_out = row['total_qty'] - row['packed_qty']
            rec = sup_saved.get(row['product_name'], {})
            # Skip items waar HO qty op 0 gestel het (✕ gedruk en gestoor)
            if rec and float(rec.get('qty_ordered') or 0) == 0:
                continue
            if sup not in suppliers:
                suppliers[sup] = []
            suppliers[sup].append({
                'product_name': row['product_name'],
                'category':     row['category'],
                'qty_needed':   qty_out,
                'qty_ordered':  rec.get('qty_ordered', qty_out) if rec else qty_out,
                'notes':        rec.get('notes', '') if rec else '',
            })
        for pname in extra_names:
            rec  = sup_saved[pname]
            info = extra_prod_info.get(pname, {'supplier': 'No Supplier', 'category': 'Salon'})
            sup_key = info['supplier']
            if sup_key not in suppliers:
                suppliers[sup_key] = []
            suppliers[sup_key].append({
                'product_name': pname,
                'category':     info['category'],
                'qty_needed':   0,
                'qty_ordered':  float(rec.get('qty_ordered') or 0),
                'notes':        rec.get('notes', '') or '',
            })

        # ── Colour palette ────────────────────────────────────────────────────
        NAVY  = RGBColor(0x15, 0x65, 0xC0)  # #1565C0
        LBLUE = RGBColor(0xE3, 0xF2, 0xFD)  # #E3F2FD
        BLUE2 = RGBColor(0x90, 0xCA, 0xF9)  # #90CAF9
        RED   = RGBColor(0xC0, 0x39, 0x2B)  # #C0392B
        DARK  = RGBColor(0x2C, 0x1A, 0x3E)  # #2C1A3E
        WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        BORDER = '90CAF9'

        # ── Helper functions ──────────────────────────────────────────────────
        def shd(cell, rgb):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            s  = OxmlElement('w:shd')
            s.set(qn('w:val'),   'clear')
            s.set(qn('w:color'), 'auto')
            s.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
            pr.append(s)

        def borders(cell, color=BORDER, sz='4', sides=('top','bottom','left','right')):
            tc = cell._tc; pr = tc.get_or_add_tcPr()
            b  = OxmlElement('w:tcBorders')
            for side in sides:
                e = OxmlElement(f'w:{side}')
                e.set(qn('w:val'),   'single')
                e.set(qn('w:sz'),    sz)
                e.set(qn('w:space'), '0')
                e.set(qn('w:color'), color)
                b.append(e)
            pr.append(b)

        def no_borders(cell):
            borders(cell, color='FFFFFF', sz='0')

        def cell_para(cell, space_before=2, space_after=2):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after  = Pt(space_after)
            return p

        def fmt(qty):
            v = float(qty) if qty is not None else 0.0
            return str(int(v)) if v == int(v) else str(v)

        def lock_tbl_width(tbl, total_cm):
            tblPr = tbl._tbl.tblPr
            for old in tblPr.findall(qn('w:tblW')):
                tblPr.remove(old)
            tw = OxmlElement('w:tblW')
            tw.set(qn('w:w'), str(int(total_cm * 1440 / 2.54)))
            tw.set(qn('w:type'), 'dxa')
            tblPr.append(tw)

        def force_col_widths(tbl, widths_cm):
            twips = [int(w * 1440 / 2.54) for w in widths_cm]
            tblGrid = tbl._tbl.tblGrid
            grid_cols = tblGrid.findall(qn('w:gridCol'))
            for i, tw in enumerate(twips):
                if i < len(grid_cols):
                    grid_cols[i].set(qn('w:w'), str(tw))
                else:
                    gc = OxmlElement('w:gridCol')
                    gc.set(qn('w:w'), str(tw))
                    tblGrid.append(gc)
            for row in tbl.rows:
                for ci, cell in enumerate(row.cells):
                    if ci >= len(widths_cm): break
                    tcPr = cell._tc.get_or_add_tcPr()
                    for old in tcPr.findall(qn('w:tcW')):
                        tcPr.remove(old)
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(twips[ci]))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)
            lock_tbl_width(tbl, sum(widths_cm))

        def page_break(doc):
            p   = doc.add_paragraph()
            run = p.add_run()
            br  = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._r.append(br)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

        # ── Document setup — Portrait A4 ──────────────────────────────────────
        doc  = Document()
        norm = doc.styles['Normal']
        norm.font.name = 'Calibri'
        norm.font.size = Pt(10)

        sec = doc.sections[0]
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)

        # Usable width = 17.4 cm
        W_CB   = 0.7
        W_PROD = 5.8
        W_CAT  = 2.2
        W_NEED = 1.8
        W_ORD  = 1.8
        W_NOTE = 5.1
        # Total: 0.7+5.8+2.2+1.8+1.8+5.1 = 17.4 cm ✓

        # ── Page header ───────────────────────────────────────────────────────
        hdr_tbl = doc.add_table(rows=1, cols=2)
        hdr_tbl.style     = 'Table Grid'
        hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_tbl.autofit   = False
        force_col_widths(hdr_tbl, [10.5, 6.9])

        # Left cell: SVS title
        c0 = hdr_tbl.rows[0].cells[0]
        shd(c0, NAVY); borders(c0)
        p0 = cell_para(c0, 5, 2)
        r0a = p0.add_run('SVS Stock')
        r0a.bold = True; r0a.font.size = Pt(18); r0a.font.color.rgb = WHITE
        r0b = p0.add_run('  ·  Bestelys / Order List')
        r0b.font.size = Pt(10); r0b.font.color.rgb = LBLUE
        p0s = c0.add_paragraph()
        p0s.paragraph_format.space_before = Pt(1)
        p0s.paragraph_format.space_after  = Pt(5)
        r0c = p0s.add_run('Silver Violet Studios – uitstaande items gegroepeer per verskaffer')
        r0c.font.size = Pt(8); r0c.font.color.rgb = BLUE2

        # Right cell: week + generated timestamp
        c1 = hdr_tbl.rows[0].cells[1]
        shd(c1, LBLUE); borders(c1)
        p1 = cell_para(c1, 5, 2)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1a = p1.add_run(f'Week: {monday} → {friday}')
        r1a.bold = True; r1a.font.size = Pt(9); r1a.font.color.rgb = NAVY
        p1s = c1.add_paragraph()
        p1s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1s.paragraph_format.space_before = Pt(2)
        p1s.paragraph_format.space_after  = Pt(5)
        r1b = p1s.add_run(f'Gegenereer: {sast_now()[:16]}')
        r1b.font.size = Pt(7.5); r1b.font.color.rgb = NAVY

        # Spacer after header
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(10)
        sp.paragraph_format.space_after  = Pt(0)

        # ── Per-supplier sections ─────────────────────────────────────────────
        first_sup = True
        for sup_name, items in suppliers.items():
            if not first_sup:
                sp2 = doc.add_paragraph()
                sp2.paragraph_format.space_before = Pt(10)
                sp2.paragraph_format.space_after  = Pt(0)
            first_sup = False

            # Supplier heading bar
            sh_tbl = doc.add_table(rows=1, cols=1)
            sh_tbl.style     = 'Table Grid'
            sh_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            sh_tbl.autofit   = False
            force_col_widths(sh_tbl, [17.4])
            sh_cell = sh_tbl.rows[0].cells[0]
            shd(sh_cell, NAVY)
            borders(sh_cell, color='1565C0')
            sh_p = cell_para(sh_cell, 4, 4)
            sh_r1 = sh_p.add_run(f'  {sup_name.upper()}')
            sh_r1.bold = True; sh_r1.font.size = Pt(10); sh_r1.font.color.rgb = WHITE
            cnt = len(items)
            sh_r2 = sh_p.add_run(f'    ({cnt} item{"s" if cnt != 1 else ""})')
            sh_r2.font.size = Pt(8); sh_r2.font.color.rgb = LBLUE

            # Product table
            tbl = doc.add_table(rows=1, cols=6)
            tbl.style     = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit   = False
            force_col_widths(tbl, [W_CB, W_PROD, W_CAT, W_NEED, W_ORD, W_NOTE])

            # Column header row
            hrow = tbl.rows[0]
            col_headers = ['☐', 'Produk / Product', 'Kategorie', 'Benodig', 'Bestel Qty', 'Notas / Remarks']
            col_aligns  = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                           WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                           WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            for ci, (htxt, halign) in enumerate(zip(col_headers, col_aligns)):
                hc = hrow.cells[ci]
                shd(hc, LBLUE); borders(hc, color=BORDER)
                hp = cell_para(hc, 2, 2)
                hp.alignment = halign
                hr = hp.add_run(htxt)
                hr.bold = True; hr.font.size = Pt(8); hr.font.color.rgb = NAVY

            # Data rows
            for ri, item in enumerate(items):
                drow = tbl.add_row()
                is_extra = item.get('qty_needed', 1) == 0
                bg = LBLUE if ri % 2 == 0 else WHITE

                # CB column
                c_cb = drow.cells[0]
                shd(c_cb, bg); borders(c_cb, color=BORDER)
                pcb = cell_para(c_cb, 3, 3)
                pcb.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pcb.add_run('☐').font.size = Pt(11)

                # Product name
                c_pr = drow.cells[1]
                shd(c_pr, bg); borders(c_pr, color=BORDER)
                ppr = cell_para(c_pr, 3, 3)
                rpr = ppr.add_run(item['product_name'])
                rpr.bold = True; rpr.font.size = Pt(9); rpr.font.color.rgb = DARK
                if is_extra:
                    rex = ppr.add_run('  +EXTRA')
                    rex.font.size = Pt(7); rex.font.color.rgb = NAVY

                # Category
                c_ca = drow.cells[2]
                shd(c_ca, bg); borders(c_ca, color=BORDER)
                pca = cell_para(c_ca, 3, 3)
                pca.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pca.add_run(item['category']).font.size = Pt(8)

                # Needed qty
                c_ne = drow.cells[3]
                shd(c_ne, bg); borders(c_ne, color=BORDER)
                pne = cell_para(c_ne, 3, 3)
                pne.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rne = pne.add_run('—' if is_extra else fmt(item['qty_needed']))
                rne.bold = True; rne.font.size = Pt(10)
                rne.font.color.rgb = DARK if is_extra else RED

                # Order qty (user-adjusted)
                c_or = drow.cells[4]
                shd(c_or, bg); borders(c_or, color=BORDER)
                por = cell_para(c_or, 3, 3)
                por.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ror = por.add_run(fmt(item['qty_ordered']))
                ror.bold = True; ror.font.size = Pt(11); ror.font.color.rgb = NAVY

                # Notes / write-in space
                c_no = drow.cells[5]
                shd(c_no, bg); borders(c_no, color=BORDER)
                pno = cell_para(c_no, 3, 3)
                note_txt = item.get('notes', '') or ''
                rno = pno.add_run(note_txt if note_txt else ' ')
                rno.font.size = Pt(8.5)
                rno.font.color.rgb = DARK if note_txt else WHITE

        # ── Footer ────────────────────────────────────────────────────────────
        fsp = doc.add_paragraph()
        fsp.paragraph_format.space_before = Pt(14)
        fsp.paragraph_format.space_after  = Pt(0)
        f_tbl = doc.add_table(rows=1, cols=2)
        f_tbl.style     = 'Table Grid'
        f_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        f_tbl.autofit   = False
        force_col_widths(f_tbl, [11.4, 6.0])
        fc0 = f_tbl.rows[0].cells[0]; fc1 = f_tbl.rows[0].cells[1]
        shd(fc0, LBLUE); no_borders(fc0)
        shd(fc1, LBLUE); no_borders(fc1)
        fp0 = cell_para(fc0, 3, 3)
        fp0.add_run('Silver Violet Studios – SVS Stock System').font.size = Pt(7.5)
        fp1 = cell_para(fc1, 3, 3)
        fp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rfp = fp1.add_run(f'Week: {monday}')
        rfp.font.size = Pt(7.5); rfp.font.color.rgb = NAVY

        # ── Serve as download ─────────────────────────────────────────────────
        import tempfile as _tmpfile, os as _os
        with _tmpfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        doc.save(tmp_path)
        with open(tmp_path, 'rb') as f:
            data = f.read()
        _os.unlink(tmp_path)

        fname = f'SVS_Bestelys_{monday}.docx'
        self.set_header('Content-Type',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.set_header('Content-Disposition', f'attachment; filename="{fname}"')
        self.write(data)


class HOSupplierOrderSaveHandler(BaseHandler):
    """Stoor verskaffer-bestelling aanpassings vanuit die Supplier Order tab."""
    @tornado.web.authenticated
    def post(self, pack_date):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        monday = week_monday(pack_date)
        sunday = week_sunday(monday)
        db = get_db()
        items_raw = db.execute("""
            SELECT oi.product_name, oi.category,
                   SUM(oi.quantity) as total_qty,
                   SUM(CASE WHEN COALESCE(oi.packed,0)=1
                            THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                            ELSE 0 END) as packed_qty
            FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            WHERE so.pack_date BETWEEN ? AND ? AND so.status != 'draft'
              AND oi.category != 'Handoeke'
            GROUP BY oi.product_name, oi.category
            HAVING SUM(oi.quantity) > SUM(CASE WHEN COALESCE(oi.packed,0)=1
                                               THEN COALESCE(NULLIF(oi.packed_qty,0), oi.quantity)
                                               ELSE 0 END)
        """, (monday, sunday)).fetchall()
        raw_map = {r['product_name']: r for r in items_raw}
        db.execute("DELETE FROM supplier_orders WHERE pack_date=?", (monday,))
        try:
            max_idx = int(self.get_argument('max_idx', '999'))
        except Exception:
            max_idx = 999
        for i in range(max_idx + 1):
            pname = self.get_argument(f'pname_{i}', None)
            if pname is None:
                continue  # Oorgeslaane ry (verwyder of gaping) — gaan voort
            raw = raw_map.get(pname)
            if raw:
                qty_needed = raw['total_qty'] - raw['packed_qty']
                try:
                    qty_ord = float(self.get_argument(f'qty_{i}', str(qty_needed)))
                except Exception:
                    qty_ord = qty_needed
                notes = self.get_argument(f'notes_{i}', '').strip()
                db.execute("""
                    INSERT INTO supplier_orders
                      (pack_date, product_name, category, qty_needed, qty_ordered, ordered, ordered_at, notes)
                    VALUES (?,?,?,?,?,1,?,?)
                """, (monday, pname, raw['category'], qty_needed, qty_ord, sast_now(), notes))
            else:
                # Extra manually-added item (not in current outstanding)
                try:
                    qty_ord = float(self.get_argument(f'qty_{i}', '0'))
                except Exception:
                    qty_ord = 0
                if qty_ord > 0:
                    notes = self.get_argument(f'notes_{i}', '').strip()
                    # cat_{i} is set by the form (e.g. 'Tints' for tint rows)
                    cat_hint = self.get_argument(f'cat_{i}', '').strip()
                    prod_info = db.execute(
                        "SELECT category FROM products WHERE LOWER(TRIM(name))=LOWER(TRIM(?))", (pname,)
                    ).fetchone()
                    cat = prod_info['category'] if prod_info else (cat_hint or 'Salon')
                    db.execute("""
                        INSERT INTO supplier_orders
                          (pack_date, product_name, category, qty_needed, qty_ordered, ordered, ordered_at, notes)
                        VALUES (?,?,?,0,?,1,?,?)
                    """, (monday, pname, cat, qty_ord, sast_now(), notes))
        db.commit()
        db.close()
        self.redirect(f'/ho/week/{monday}#tab-supplier')


class AdminProductsHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        active_products   = db.execute(
            "SELECT * FROM products WHERE active=1 ORDER BY category, name"
        ).fetchall()
        inactive_products = db.execute(
            "SELECT * FROM products WHERE active=0 ORDER BY category, name"
        ).fetchall()
        db.close()
        self.render('ho/products.html', user=u,
                    products=[dict(p) for p in active_products],
                    inactive_products=[dict(p) for p in inactive_products])

    @tornado.web.authenticated
    def post(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        name     = self.get_argument('name', '').strip()
        category = self.get_argument('category', 'Salon').strip()
        supplier = self.get_argument('supplier', '').strip()
        if not name:
            self.redirect('/ho/products'); return
        db = get_db()
        # Vermy duplikate — as die produk al bestaan, aktiveer dit net weer
        existing = db.execute(
            "SELECT id FROM products WHERE name=? AND category=?", (name, category)
        ).fetchone()
        if existing:
            db.execute("UPDATE products SET active=1, supplier=? WHERE id=?",
                       (supplier, existing['id']))
        else:
            db.execute(
                "INSERT INTO products (name, category, supplier, active) VALUES (?,?,?,1)",
                (name, category, supplier)
            )
        db.commit()
        db.close()
        self.redirect('/ho/products')

class AdminProductToggleHandler(BaseHandler):
    """Aktiveer of deaktiveer 'n produk (stel active=0/1)."""
    @tornado.web.authenticated
    def post(self, product_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        action = self.get_argument('action', 'deactivate')
        new_active = 0 if action == 'deactivate' else 1
        db = get_db()
        db.execute("UPDATE products SET active=? WHERE id=?", (new_active, product_id))
        db.commit()
        db.close()
        self.redirect('/ho/products')

class AdminProductEditHandler(BaseHandler):
    """Verander 'n produk se naam, kategorie en/of verskaffer."""
    @tornado.web.authenticated
    def post(self, product_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        name     = self.get_argument('name', '').strip()
        category = self.get_argument('category', '').strip()
        supplier = self.get_argument('supplier', '').strip()
        if not name or not category:
            self.redirect('/ho/products'); return
        db = get_db()
        db.execute(
            "UPDATE products SET name=?, category=?, supplier=? WHERE id=?",
            (name, category, supplier, product_id)
        )
        db.commit()
        db.close()
        self.redirect('/ho/products')

class AdminProductDeleteHandler(BaseHandler):
    """Deaktiveer 'n produk (stel active=0) — verwyder NIE die ry nie sodat dit nie herlaai word nie."""
    @tornado.web.authenticated
    def post(self, product_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        db.execute("UPDATE products SET active=0 WHERE id=?", (product_id,))
        db.commit()
        db.close()
        self.redirect('/ho/products')

class AdminProductCleanupHandler(BaseHandler):
    """Verwyder duplikate — behou die een met 'n verskaffer; as almal leeg is, behou laagste ID."""
    @tornado.web.authenticated
    def post(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        # Per name+category group:
        #   - If any row has a supplier → keep the lowest id that has a supplier
        #   - If none have a supplier   → keep the lowest id overall
        db.execute("""
            DELETE FROM products
            WHERE id NOT IN (
                SELECT COALESCE(
                    MIN(CASE WHEN supplier IS NOT NULL AND TRIM(supplier) != ''
                             THEN id END),
                    MIN(id)
                )
                FROM products
                GROUP BY LOWER(TRIM(name)), LOWER(TRIM(category))
            )
        """)
        db.commit()
        db.close()
        self.redirect('/ho/products')

# ─── Order Preview / Confirm / Delete Item ────────────────────────────────────

class OrderPreviewHandler(BaseHandler):
    """Oorsig-bladsy: wys items met trash-knoppies voor finale indiening."""
    @tornado.web.authenticated
    def get(self, order_id):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND salon_id=? AND status='draft'",
            (order_id, u['salon_id'])
        ).fetchone()
        if not order:
            db.close(); self.redirect('/salon/dashboard'); return
        items = db.execute(
            "SELECT * FROM order_items WHERE order_id=? ORDER BY category, product_name",
            (order_id,)
        ).fetchall()
        db.close()
        self.render('salon/order_preview.html', user=u,
                    order=dict(order), items=[dict(i) for i in items])

class OrderConfirmHandler(BaseHandler):
    """Bevestig indiening: stoor notas per item, verander status na 'submitted'."""
    @tornado.web.authenticated
    def post(self, order_id):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND salon_id=? AND status='draft'",
            (order_id, u['salon_id'])
        ).fetchone()
        if not order:
            db.close(); self.redirect('/salon/dashboard'); return
        items = db.execute("SELECT * FROM order_items WHERE order_id=?", (order_id,)).fetchall()
        # Geld as leeg as geen items met qty > 0 (handoeke met 0 tel nie as "item" nie)
        real_items = [i for i in items if i['category'] != 'Handoeke' or float(i['quantity']) > 0]
        if len(real_items) == 0:
            db.close(); self.redirect(f'/salon/order/preview/{order_id}'); return
        # Stoor nota en hoeveelheid per item
        for item in items:
            note = self.get_argument(f'note_{item["id"]}', '').strip()
            try:
                qty = float(self.get_argument(f'qty_{item["id"]}', str(item['quantity'])))
                # Handoeke mag qty=0 hê; ander items minimum 1
                if item['category'] != 'Handoeke' and qty < 1:
                    qty = 1
            except Exception:
                qty = float(item['quantity'])
            db.execute("UPDATE order_items SET note=?, quantity=? WHERE id=?",
                       (note, qty, item['id']))
        now = sast_now()
        db.execute(
            "UPDATE stock_orders SET status='submitted', submitted_at=?, pack_date=? WHERE id=?",
            (now, week_monday(), order_id)
        )
        db.commit()
        db.close()
        self.redirect(f'/salon/order/{order_id}')

class OrderDeleteItemHandler(BaseHandler):
    """Verwyder 'n enkele item uit 'n konsep-bestelling."""
    @tornado.web.authenticated
    def post(self, item_id):
        u = self.current_user
        db = get_db()
        row = db.execute("""
            SELECT oi.id, oi.order_id, o.status, o.salon_id
            FROM order_items oi JOIN stock_orders o ON oi.order_id = o.id
            WHERE oi.id=?
        """, (item_id,)).fetchone()
        order_id = None
        if row and row['salon_id'] == u['salon_id'] and row['status'] == 'draft':
            order_id = row['order_id']
            db.execute("DELETE FROM order_items WHERE id=?", (item_id,))
            db.commit()
        db.close()
        if order_id:
            self.redirect(f'/salon/order/preview/{order_id}')
        else:
            self.redirect('/salon/dashboard')

# ─── Aflewering ───────────────────────────────────────────────────────────────

TOWEL_TYPES = ['Pink','Blue','Green','Maroon','Grey','White','Purple','Black','FC (Frailcare)','Perm']

class DeliveryHandler(BaseHandler):
    """Wys aflewerings-verifikasie bladsy en verwerk die teken-af."""
    @tornado.web.authenticated
    def get(self, order_id):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND salon_id=? AND status IN ('submitted','packed')",
            (order_id, u['salon_id'])
        ).fetchone()
        if not order:
            db.close(); self.redirect('/salon/dashboard'); return
        items = db.execute("""
            SELECT *,
              CASE WHEN COALESCE(packed, 0) = 0 THEN 0
                   WHEN COALESCE(packed_qty, 0) = 0 THEN quantity
                   ELSE packed_qty
              END as eff_packed_qty
            FROM order_items
            WHERE order_id=? AND category != 'Handoeke'
            ORDER BY
              CASE WHEN COALESCE(packed, 0) = 0 THEN 1 ELSE 0 END,
              category, product_name
        """, (order_id,)).fetchall()
        # Bestelde handdoek-hoeveelhede vir hierdie order
        ordered_towels = {}
        for r in db.execute(
            "SELECT product_name, quantity FROM order_items WHERE order_id=? AND category='Handoeke'",
            (order_id,)
        ).fetchall():
            ordered_towels[r['product_name']] = float(r['quantity'])

        # Laaste opgeteldes per handoek-tipe (as wenk as niks bestel is nie)
        last_rows = db.execute("""
            SELECT towel_type, collected FROM towel_logs
            WHERE salon_id=? AND collected > 0
            ORDER BY created_at DESC
        """, (u['salon_id'],)).fetchall()
        last_collected = {}
        for r in last_rows:
            if r['towel_type'] not in last_collected:
                last_collected[r['towel_type']] = r['collected']
        db.close()
        self.render('salon/delivery.html', user=u, order=dict(order),
                    items=[dict(i) for i in items],
                    towel_types=TOWEL_TYPES,
                    ordered_towels=ordered_towels,
                    last_collected=last_collected)

    @tornado.web.authenticated
    def post(self, order_id):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND salon_id=? AND status IN ('submitted','packed')",
            (order_id, u['salon_id'])
        ).fetchone()
        if not order:
            db.close(); self.redirect('/salon/dashboard'); return

        # Stoor afleweringsdata per item
        items = db.execute("SELECT * FROM order_items WHERE order_id=?", (order_id,)).fetchall()
        for item in items:
            iid = item['id']
            ticked = self.get_argument(f'recv_{iid}', '') == 'on'
            if ticked:
                try:
                    del_qty = float(self.get_argument(f'dqty_{iid}', str(item['quantity'])))
                except Exception:
                    del_qty = float(item['quantity'])
            else:
                del_qty = 0.0
            del_note = self.get_argument(f'dnote_{iid}', '').strip()
            db.execute(
                "UPDATE order_items SET delivered_qty=?, delivery_note=? WHERE id=?",
                (del_qty, del_note, iid)
            )

        # Stoor handoek-data (afgelewer + opgetel per tipe)
        pack_date = order['pack_date']
        db.execute("DELETE FROM towel_logs WHERE order_id=?", (order_id,))
        for i, t in enumerate(TOWEL_TYPES):
            try:
                sent = int(self.get_argument(f'tsent_{i}', '0') or 0)
            except Exception:
                sent = 0
            try:
                coll = int(self.get_argument(f'tcoll_{i}', '0') or 0)
            except Exception:
                coll = 0
            if sent > 0 or coll > 0:
                db.execute(
                    "INSERT INTO towel_logs (order_id, salon_id, week_date, towel_type, sent, collected) "
                    "VALUES (?,?,?,?,?,?)",
                    (order_id, u['salon_id'], pack_date, t, sent, coll)
                )

        # Extra items added during delivery
        for i in range(20):
            extra_name = self.get_argument(f'extra_name_{i}', '').strip()
            extra_cat  = self.get_argument(f'extra_cat_{i}', 'Salon').strip()
            try:
                extra_qty = float(self.get_argument(f'extra_qty_{i}', '0') or 0)
            except Exception:
                extra_qty = 0.0
            if extra_name and extra_qty > 0:
                db.execute("""
                    INSERT INTO order_items
                      (order_id, product_name, category, quantity, packed, packed_qty,
                       delivered_qty, note, is_custom, ho_added)
                    VALUES (?,?,?,?,1,?,?,?,1,0)
                """, (order_id, extra_name, extra_cat, extra_qty, extra_qty, extra_qty, ''))

        signed_by = self.get_argument('signed_by', '').strip()
        db.execute(
            "UPDATE stock_orders SET status='delivered', delivered_at=datetime('now'), signed_by=? WHERE id=?",
            (signed_by, order_id)
        )
        db.commit()
        db.close()
        self.redirect('/salon/dashboard')

# ─── Verander ingediene bestelling (terug na konsep) ─────────────────────────

class OrderRevertHandler(BaseHandler):
    """Verander 'n ingediene (nog nie gepakte) bestelling terug na konsep sodat dit gewysig kan word."""
    @tornado.web.authenticated
    def post(self, order_id):
        u = self.current_user
        db = get_db()
        if u['role'] == 'ho_admin':
            # HO kan enige ingediende bestelling terugkeer (nie gepakte nie)
            order = db.execute(
                "SELECT * FROM stock_orders WHERE id=? AND status='submitted'",
                (order_id,)
            ).fetchone()
            if order:
                db.execute(
                    "UPDATE stock_orders SET status='draft', submitted_at=NULL WHERE id=?",
                    (order_id,)
                )
                db.commit()
            db.close()
            self.redirect(f'/ho/order/{order_id}/edit')
        else:
            # Salon kan slegs hul eie bestelling terugkeer
            order = db.execute(
                "SELECT * FROM stock_orders WHERE id=? AND salon_id=? AND status='submitted'",
                (order_id, u['salon_id'])
            ).fetchone()
            if order:
                db.execute(
                    "UPDATE stock_orders SET status='draft', submitted_at=NULL WHERE id=?",
                    (order_id,)
                )
                db.commit()
            db.close()
            self.redirect('/salon/order/new')


# ─── Staan-alleen Handoeke ────────────────────────────────────────────────────

class SalonTowelHandler(BaseHandler):
    """Rekord handoeke-aflewering/kolleksie onafhanklik van 'n stock bestelling."""
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        db = get_db()
        # Aktiewe bestelling se handoeke (Opgetel Vuil — vaste bedrag van die bestelling)
        active_order_towels = {}
        active_order = db.execute("""
            SELECT id FROM stock_orders
            WHERE salon_id=? AND status IN ('submitted','packed')
            ORDER BY submitted_at DESC LIMIT 1
        """, (u['salon_id'],)).fetchone()
        if active_order:
            for r in db.execute("""
                SELECT product_name, quantity FROM order_items
                WHERE order_id=? AND category='Handoeke'
            """, (active_order['id'],)).fetchall():
                active_order_towels[r['product_name']] = float(r['quantity'])
        # Mees onlangse towel_logs rekords per tipe (vir voor-invul verwysing)
        last_towels = {}
        for r in db.execute("""
            SELECT towel_type, sent, collected, COALESCE(next_qty, 0) as next_qty
            FROM towel_logs WHERE salon_id=? ORDER BY created_at DESC
        """, (u['salon_id'],)).fetchall():
            if r['towel_type'] not in last_towels:
                last_towels[r['towel_type']] = {
                    'sent':      float(r['sent'] or 0),
                    'collected': float(r['collected'] or 0),
                    'next_qty':  float(r['next_qty'] or 0)
                }
        db.close()
        today = datetime.date.today().isoformat()
        self.render('salon/towels.html', user=u,
                    towel_types=TOWEL_TYPES,
                    active_order_towels=active_order_towels,
                    last_towels=last_towels,
                    today=today)

    @tornado.web.authenticated
    def post(self):
        u = self.current_user
        if u['role'] == 'ho_admin':
            self.redirect('/ho/dashboard'); return
        db = get_db()
        week_date = self.get_argument('week_date', datetime.date.today().isoformat())
        # Verwyder bestaande staan-alleen rekords vir hierdie salone+datum
        db.execute(
            "DELETE FROM towel_logs WHERE salon_id=? AND order_id IS NULL AND week_date=?",
            (u['salon_id'], week_date)
        )
        for i, t in enumerate(TOWEL_TYPES):
            try:
                sent = float(self.get_argument(f'tsent_{i}', '0') or 0)   # Terug Ontvang Skoon
            except Exception:
                sent = 0.0
            try:
                coll = float(self.get_argument(f'tcoll_{i}', '0') or 0)   # Opgetel Vuil (van bestelling)
            except Exception:
                coll = 0.0
            try:
                next_q = float(self.get_argument(f'tnext_{i}', '0') or 0) # Gestuur vir volgende week
            except Exception:
                next_q = 0.0
            if sent > 0 or coll > 0 or next_q > 0:
                db.execute(
                    "INSERT INTO towel_logs (order_id, salon_id, week_date, towel_type, sent, collected, next_qty) "
                    "VALUES (NULL,?,?,?,?,?,?)",
                    (u['salon_id'], week_date, t, sent, coll, next_q)
                )
        db.commit()
        db.close()
        self.redirect('/salon/dashboard')


# ─── HO Stock Take ────────────────────────────────────────────────────────────

class HOStockTakeListHandler(BaseHandler):
    """Lys van alle stock takes."""
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        takes_raw = db.execute("""
            SELECT st.*, COUNT(sti.id) as item_count
            FROM stock_takes st
            LEFT JOIN stock_take_items sti ON sti.stock_take_id = st.id
            GROUP BY st.id
            ORDER BY st.taken_at DESC LIMIT 100
        """).fetchall()
        db.close()
        self.render('ho/stocktake_list.html', user=u, takes=[dict(t) for t in takes_raw])


class HOStockTakeNewHandler(BaseHandler):
    """Nuwe stock take skerm — Tel & Stoor."""
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        salon_products    = [dict(p) for p in db.execute(
            "SELECT * FROM products WHERE category='Salon' AND active=1 ORDER BY name"
        ).fetchall()]
        cleaning_products = [dict(p) for p in db.execute(
            "SELECT * FROM products WHERE category='Cleaning' AND active=1 ORDER BY name"
        ).fetchall()]
        perm_products     = [dict(p) for p in db.execute(
            "SELECT * FROM products WHERE category='Perms' AND active=1 ORDER BY name"
        ).fetchall()]
        retail_products   = [dict(p) for p in db.execute(
            "SELECT * FROM products WHERE category='Retail' AND active=1 ORDER BY name"
        ).fetchall()]
        db.close()
        today = datetime.date.today().isoformat()
        self.render('ho/stocktake_new.html', user=u,
                    salon_products=salon_products,
                    cleaning_products=cleaning_products,
                    perm_products=perm_products,
                    retail_products=retail_products,
                    today=today)

    @tornado.web.authenticated
    def post(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        taken_at = self.get_argument('taken_at', datetime.date.today().isoformat())
        taken_by = self.get_argument('taken_by', u['name']).strip()
        notes    = self.get_argument('notes', '').strip()
        db = get_db()
        cur = db.execute(
            "INSERT INTO stock_takes (taken_at, taken_by, notes) VALUES (?,?,?)",
            (taken_at, taken_by, notes)
        )
        stock_take_id = cur.lastrowid
        all_products = db.execute(
            "SELECT * FROM products WHERE category IN ('Salon','Cleaning','Perms','Retail') AND active=1"
        ).fetchall()
        for p in all_products:
            try:
                qty = float(self.get_argument(f'qty_{p["id"]}', '0') or 0)
            except Exception:
                qty = 0.0
            if qty > 0:
                db.execute(
                    "INSERT INTO stock_take_items (stock_take_id, product_name, category, quantity) "
                    "VALUES (?,?,?,?)",
                    (stock_take_id, p['name'], p['category'], qty)
                )
        db.commit()
        db.close()
        self.redirect(f'/ho/stocktake/{stock_take_id}')


class HOStockTakeViewHandler(BaseHandler):
    """Wys 'n spesifieke stock take."""
    @tornado.web.authenticated
    def get(self, take_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        take = db.execute("SELECT * FROM stock_takes WHERE id=?", (take_id,)).fetchone()
        if not take:
            db.close(); self.redirect('/ho/stocktake'); return
        items = db.execute(
            "SELECT * FROM stock_take_items WHERE stock_take_id=? ORDER BY CASE category WHEN 'Salon' THEN 1 WHEN 'Cleaning' THEN 2 WHEN 'Perms' THEN 3 WHEN 'Retail' THEN 4 ELSE 5 END, product_name",
            (take_id,)
        ).fetchall()
        db.close()
        by_cat = {}
        for item in items:
            cat = item['category']
            if cat not in by_cat:
                by_cat[cat] = []
            by_cat[cat].append(dict(item))
        self.render('ho/stocktake_view.html', user=u,
                    take=dict(take), by_cat=by_cat,
                    categories=['Salon', 'Cleaning', 'Perms', 'Retail'])


class HOStockTakeXlsxHandler(BaseHandler):
    """Genereer Excel-uitvoer van 'n stock take."""
    @tornado.web.authenticated
    def get(self, take_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
            import io
        except ImportError:
            self.write('<html><body style="font-family:Arial;padding:2rem">'
                       '<h2>openpyxl nie geïnstalleer nie</h2>'
                       '<p>Hardloop: <code>pip install openpyxl</code></p>'
                       '<a href="/ho/stocktake">Terug</a></body></html>')
            return
        db = get_db()
        take  = db.execute("SELECT * FROM stock_takes WHERE id=?", (take_id,)).fetchone()
        if not take:
            db.close(); self.redirect('/ho/stocktake'); return
        items = db.execute(
            "SELECT * FROM stock_take_items WHERE stock_take_id=? ORDER BY CASE category WHEN 'Salon' THEN 1 WHEN 'Cleaning' THEN 2 WHEN 'Perms' THEN 3 WHEN 'Retail' THEN 4 ELSE 5 END, product_name",
            (take_id,)
        ).fetchall()
        db.close()
        take = dict(take)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Stock Take"

        # Style helpers
        def fill(hex_color):
            return PatternFill("solid", fgColor=hex_color)
        def font(bold=False, size=11, color="000000"):
            return Font(bold=bold, size=size, color=color)
        def align(h='left', v='center'):
            return Alignment(horizontal=h, vertical=v)

        PURPLE = "3D1F5C"
        LILAC  = "7B3FA0"
        SILVER = "F0ECF6"
        WHITE  = "FFFFFF"
        GREY   = "7A6B8A"

        # ── Row 1: Title ──
        ws.merge_cells('A1:C1')
        c = ws['A1']
        c.value     = f"SVS Stock Take – {take['taken_at']}"
        c.font      = font(bold=True, size=14, color=WHITE)
        c.fill      = fill(PURPLE)
        c.alignment = align('center')
        ws.row_dimensions[1].height = 30

        # ── Row 2: Info ──
        ws.merge_cells('A2:C2')
        c = ws['A2']
        parts = []
        if take.get('taken_by'):
            parts.append(f"Geneem deur: {take['taken_by']}")
        if take.get('notes'):
            parts.append(f"Notas: {take['notes']}")
        c.value     = '  |  '.join(parts)
        c.font      = font(size=10, color=GREY)
        c.alignment = align('left')
        ws.row_dimensions[2].height = 18

        # ── Row 3: Column headers ──
        for ci, (lbl, al) in enumerate([('Kategorie','center'),('Produk','left'),('Hoeveelheid','center')], 1):
            c = ws.cell(row=3, column=ci, value=lbl)
            c.font      = font(bold=True, size=11, color=WHITE)
            c.fill      = fill(LILAC)
            c.alignment = align(al)
        ws.row_dimensions[3].height = 22

        row = 4
        current_cat = None
        bg_toggle   = True
        for item in items:
            item = dict(item)
            if item['category'] != current_cat:
                current_cat = item['category']
                ws.merge_cells(f'A{row}:C{row}')
                c = ws.cell(row=row, column=1, value=f"  {current_cat.upper()}")
                c.font      = font(bold=True, size=11, color=WHITE)
                c.fill      = fill(PURPLE)
                c.alignment = align('left')
                ws.row_dimensions[row].height = 20
                row += 1
                bg_toggle = True

            bg = SILVER if bg_toggle else "FFFFFF"
            qty = item['quantity']
            qty_val = int(qty) if qty == int(qty) else qty

            ws.cell(row=row, column=1, value=item['category']).fill = fill(bg)
            ws['A' + str(row)].font      = font(size=10, color=GREY)
            ws['A' + str(row)].alignment = align('left')

            ws.cell(row=row, column=2, value=item['product_name']).fill = fill(bg)
            ws['B' + str(row)].font      = font(size=10)
            ws['B' + str(row)].alignment = align('left')

            ws.cell(row=row, column=3, value=qty_val).fill = fill(bg)
            ws['C' + str(row)].font      = font(bold=True, size=11, color=PURPLE)
            ws['C' + str(row)].alignment = align('center')

            ws.row_dimensions[row].height = 18
            row      += 1
            bg_toggle = not bg_toggle

        ws.column_dimensions['A'].width = 14
        ws.column_dimensions['B'].width = 36
        ws.column_dimensions['C'].width = 14

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        data = buf.read()
        date_str = take['taken_at'].replace('-', '')
        fname = f"stock_take_{date_str}.xlsx"
        self.set_header('Content-Type',
                        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        self.set_header('Content-Disposition', f'attachment; filename="{fname}"')
        self.write(data)


class HOStockTakePdfHandler(BaseHandler):
    """Genereer PDF-uitvoer van 'n stock take."""
    @tornado.web.authenticated
    def get(self, take_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                             Paragraph, Spacer)
            from reportlab.lib.styles import ParagraphStyle
            import io
        except ImportError:
            self.write('<html><body style="font-family:Arial;padding:2rem">'
                       '<h2>reportlab nie geïnstalleer nie</h2>'
                       '<p>Hardloop: <code>pip install reportlab</code></p>'
                       '<a href="/ho/stocktake">Terug</a></body></html>')
            return

        db = get_db()
        take  = db.execute("SELECT * FROM stock_takes WHERE id=?", (take_id,)).fetchone()
        if not take:
            db.close(); self.redirect('/ho/stocktake'); return
        items = db.execute(
            "SELECT * FROM stock_take_items WHERE stock_take_id=? ORDER BY CASE category WHEN 'Salon' THEN 1 WHEN 'Cleaning' THEN 2 WHEN 'Perms' THEN 3 WHEN 'Retail' THEN 4 ELSE 5 END, product_name",
            (take_id,)
        ).fetchall()
        db.close()
        take  = dict(take)
        items = [dict(i) for i in items]

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=1.5*cm, rightMargin=1.5*cm,
                                topMargin=2*cm, bottomMargin=2*cm)

        PURPLE = colors.HexColor('#3D1F5C')
        LILAC  = colors.HexColor('#7B3FA0')
        SILVER = colors.HexColor('#F0ECF6')
        GREY   = colors.HexColor('#7A6B8A')
        WHITE  = colors.white

        def ps(name, **kw):
            return ParagraphStyle(name, **kw)

        story = []

        # ── Title block ──
        title_style = ps('T', fontName='Helvetica-Bold', fontSize=14,
                         textColor=WHITE, spaceAfter=0)
        story.append(Table(
            [[Paragraph(f"SVS Stock Take – {take['taken_at']}", title_style)]],
            colWidths=[17*cm],
            style=TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), PURPLE),
                ('TOPPADDING',    (0,0), (-1,-1), 10),
                ('BOTTOMPADDING', (0,0), (-1,-1), 10),
                ('LEFTPADDING',   (0,0), (-1,-1), 14),
            ])
        ))

        # ── Info line ──
        parts = []
        if take.get('taken_by'):
            parts.append(f"Geneem deur: {take['taken_by']}")
        if take.get('notes'):
            parts.append(f"Notas: {take['notes']}")
        if parts:
            story.append(Spacer(1, 0.2*cm))
            story.append(Paragraph(' | '.join(parts),
                                   ps('I', fontSize=9, textColor=GREY)))
        story.append(Spacer(1, 0.5*cm))

        # ── Build table ──
        hdr_style  = ps('H', fontName='Helvetica-Bold', fontSize=10, textColor=WHITE)
        prod_style = ps('P', fontSize=10)
        qty_style  = ps('Q', fontName='Helvetica-Bold', fontSize=11,
                        alignment=1, textColor=PURPLE)  # alignment 1 = CENTER
        cat_style  = ps('C', fontName='Helvetica-Bold', fontSize=10, textColor=WHITE)

        table_data  = [[Paragraph('Produk', hdr_style),
                        Paragraph('Hoeveelheid', hdr_style)]]
        style_cmds  = [
            ('BACKGROUND',    (0,0), (-1,0), LILAC),
            ('ALIGN',         (1,0), (1,-1), 'CENTER'),
            ('GRID',          (0,0), (-1,-1), 0.5, colors.HexColor('#D5C8E8')),
            ('TOPPADDING',    (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('LEFTPADDING',   (0,0), (-1,-1), 8),
            ('RIGHTPADDING',  (0,0), (-1,-1), 8),
        ]

        data_row   = 1
        current_cat = None
        bg_toggle   = True
        for item in items:
            if item['category'] != current_cat:
                current_cat = item['category']
                table_data.append([Paragraph(current_cat.upper(), cat_style), ''])
                style_cmds.append(('BACKGROUND', (0, data_row), (-1, data_row), PURPLE))
                style_cmds.append(('SPAN',       (0, data_row), (-1, data_row)))
                data_row  += 1
                bg_toggle  = True

            qty     = item['quantity']
            qty_str = str(int(qty)) if qty == int(qty) else str(qty)
            if bg_toggle:
                style_cmds.append(('BACKGROUND', (0, data_row), (-1, data_row), SILVER))

            table_data.append([
                Paragraph(item['product_name'], prod_style),
                Paragraph(qty_str, qty_style),
            ])
            data_row  += 1
            bg_toggle  = not bg_toggle

        if len(table_data) > 1:
            story.append(Table(table_data, colWidths=[13*cm, 4*cm],
                               style=TableStyle(style_cmds)))

        # ── Footer ──
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph(
            f"Gegenereer: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
            ps('F', fontSize=8, textColor=GREY)
        ))

        doc.build(story)
        buf.seek(0)
        data = buf.read()
        date_str = take['taken_at'].replace('-', '')
        fname = f"stock_take_{date_str}.pdf"
        self.set_header('Content-Type', 'application/pdf')
        self.set_header('Content-Disposition', f'attachment; filename="{fname}"')
        self.write(data)


# ─── Per-item mark as delivered ───────────────────────────────────────────────

class HOMarkItemDeliveredHandler(BaseHandler):
    """Merk 'n enkele item as volledig afgelewer (dismiss outstanding)."""
    @tornado.web.authenticated
    def post(self, item_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        item = db.execute("""
            SELECT oi.*, so.id as order_id FROM order_items oi
            JOIN stock_orders so ON oi.order_id = so.id
            WHERE oi.id=? AND so.status='delivered'
        """, (item_id,)).fetchone()
        if item:
            db.execute("UPDATE order_items SET delivered_qty=quantity WHERE id=?", (item_id,))
            db.commit()
            order_id = item['order_id']
            db.close()
            self.redirect(f'/ho/order/{order_id}')
        else:
            db.close()
            self.redirect('/ho/dashboard')


# ─── Billing / Month Handlers ─────────────────────────────────────────────────

class HOOrderDeleteHandler(BaseHandler):
    @tornado.web.authenticated
    def post(self, order_id):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        order = db.execute(
            "SELECT * FROM stock_orders WHERE id=? AND status IN ('submitted','packed','delivered')",
            (order_id,)
        ).fetchone()
        if order:
            pack_date = order['pack_date']
            db.execute("DELETE FROM order_items WHERE order_id=?", (order_id,))
            db.execute("DELETE FROM stock_orders WHERE id=?", (order_id,))
            db.commit()
            db.close()
            self.redirect(f'/ho/week/{pack_date}')
        else:
            db.close()
            self.redirect('/ho/dashboard')


class HOMonthListHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        rows = db.execute("""
            SELECT strftime('%Y', delivered_at) as year,
                   strftime('%m', delivered_at) as month,
                   COUNT(DISTINCT salon_id) as salon_count,
                   COUNT(*) as order_count
            FROM stock_orders
            WHERE status='delivered' AND delivered_at IS NOT NULL
            GROUP BY year, month
            ORDER BY year DESC, month DESC
        """).fetchall()
        db.close()
        months = []
        for r in rows:
            y, m = int(r['year']), int(r['month'])
            months.append({
                'year': y, 'month': m,
                'label': f"{MONTHS_FULL[m-1]} {y}",
                'salon_count': r['salon_count'],
                'order_count': r['order_count'],
            })
        self.render('ho/months.html', user=u, months=months)


class HOMonthHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self, year, month):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        year, month = int(year), int(month)
        month_str = f"{year}-{month:02d}"
        orders = db.execute("""
            SELECT o.id, o.salon_id, o.delivered_at, o.signed_by, o.pack_date,
                   s.code as salon_code, s.name as salon_name
            FROM stock_orders o
            JOIN salons s ON o.salon_id = s.id
            WHERE o.status='delivered'
              AND strftime('%Y-%m', o.delivered_at) = ?
            ORDER BY s.code, o.delivered_at
        """, (month_str,)).fetchall()
        salon_summaries = {}
        for o in orders:
            code = o['salon_code']
            if code not in salon_summaries:
                salon_summaries[code] = {
                    'salon_code': code,
                    'salon_name': o['salon_name'],
                    'order_count': 0,
                    'total_qty': 0,
                    'last_delivery': None,
                }
            salon_summaries[code]['order_count'] += 1
            if o['delivered_at']:
                salon_summaries[code]['last_delivery'] = o['delivered_at'][:10]
            items = db.execute("""
                SELECT COALESCE(delivered_qty, quantity) as del_qty
                FROM order_items
                WHERE order_id=? AND category != 'Handoeke'
                  AND COALESCE(delivered_qty, 0) > 0
            """, (o['id'],)).fetchall()
            salon_summaries[code]['total_qty'] += sum(float(i['del_qty']) for i in items)
        db.close()
        salons = sorted(salon_summaries.values(), key=lambda x: x['salon_code'])
        month_label = f"{MONTHS_FULL[month-1]} {year}"
        self.render('ho/month.html', user=u, salons=salons,
                    year=year, month=month, month_label=month_label)


class HOMonthSalonHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self, year, month, salon_code):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        year, month = int(year), int(month)
        month_str = f"{year}-{month:02d}"
        orders = db.execute("""
            SELECT o.id, o.delivered_at, o.signed_by, o.pack_date,
                   s.code as salon_code, s.name as salon_name
            FROM stock_orders o
            JOIN salons s ON o.salon_id = s.id
            WHERE o.status='delivered'
              AND strftime('%Y-%m', o.delivered_at) = ?
              AND s.code = ?
            ORDER BY o.delivered_at
        """, (month_str, salon_code)).fetchall()
        if not orders:
            self.redirect(f'/ho/month/{year}/{month:02d}'); return
        salon_name = orders[0]['salon_name']
        # Delivery dates in order (short label e.g. "14 Jul")
        dates = []
        date_labels = []
        date_info = []  # {date, label, signed_by}
        for o in orders:
            d = (o['delivered_at'] or o['pack_date'])[:10]
            if d not in dates:
                dates.append(d)
                dt = datetime.date.fromisoformat(d)
                lbl = f"{dt.day} {MAANDE[dt.month-1]}"
                date_labels.append(lbl)
                date_info.append({'date': d, 'label': lbl, 'signed_by': o['signed_by'] or ''})

        # Aggregate qty per (category, product) per date
        product_data = {}  # (cat, name) -> {date: qty}
        for o in orders:
            d = (o['delivered_at'] or o['pack_date'])[:10]
            rows = db.execute("""
                SELECT product_name, category,
                       COALESCE(delivered_qty, quantity) as del_qty
                FROM order_items
                WHERE order_id=? AND category != 'Handoeke'
                  AND COALESCE(delivered_qty, 0) > 0
                ORDER BY category, product_name
            """, (o['id'],)).fetchall()
            for row in rows:
                key = (row['category'], row['product_name'])
                if key not in product_data:
                    product_data[key] = {}
                product_data[key][d] = product_data[key].get(d, 0) + float(row['del_qty'])

        cat_order = ['Salon', 'Cleaning', 'Perms', 'Tints', 'Retail']
        categories = []
        grand_total = 0
        for cat in cat_order:
            cat_rows = []
            for (c, name), date_qtys in sorted(product_data.items()):
                if c != cat: continue
                qtys = [date_qtys.get(d, 0) for d in dates]
                total = sum(qtys)
                qtys_fmt = [int(q) if q == int(q) else q for q in qtys]
                total_fmt = int(total) if total == int(total) else total
                cat_rows.append({'product_name': name, 'qtys': qtys_fmt, 'total': total_fmt})
                grand_total += total
            if cat_rows:
                cat_total = sum(r['total'] for r in cat_rows)
                categories.append({
                    'name': cat,
                    'rows': cat_rows,
                    'total': int(cat_total) if cat_total == int(cat_total) else cat_total,
                })

        grand_total_v = int(grand_total) if grand_total == int(grand_total) else grand_total
        month_label = f"{MONTHS_FULL[month-1]} {year}"
        db.close()
        self.render('ho/month_salon.html', user=u,
                    salon_code=salon_code, salon_name=salon_name,
                    dates=dates, date_labels=date_labels, date_info=date_info,
                    categories=categories, grand_total=grand_total_v,
                    orders=[dict(o) for o in orders],
                    year=year, month=month, month_label=month_label)


class HOMonthCombinedHandler(BaseHandler):
    @tornado.web.authenticated
    def get(self, year, month):
        u = self.current_user
        if u['role'] != 'ho_admin':
            self.redirect('/salon/dashboard'); return
        db = get_db()
        year, month = int(year), int(month)
        month_str = f"{year}-{month:02d}"
        # All delivered order_items for this month (excl. towels)
        rows = db.execute("""
            SELECT oi.product_name, oi.category,
                   COALESCE(oi.delivered_qty, oi.quantity) as del_qty,
                   s.code as salon_code
            FROM order_items oi
            JOIN stock_orders o ON oi.order_id = o.id
            JOIN salons s ON o.salon_id = s.id
            WHERE o.status = 'delivered'
              AND strftime('%Y-%m', o.delivered_at) = ?
              AND oi.category != 'Handoeke'
              AND COALESCE(oi.delivered_qty, 0) > 0
            ORDER BY oi.category, oi.product_name, s.code
        """, (month_str,)).fetchall()
        db.close()

        from collections import OrderedDict
        agg = OrderedDict()
        for row in rows:
            key = (row['category'], row['product_name'])
            if key not in agg:
                agg[key] = {
                    'category': row['category'],
                    'product_name': row['product_name'],
                    'total_qty': 0,
                    'salons': []
                }
            qty = float(row['del_qty'])
            agg[key]['total_qty'] += qty
            existing = next((s for s in agg[key]['salons'] if s['code'] == row['salon_code']), None)
            if existing:
                existing['qty'] += qty
            else:
                agg[key]['salons'].append({'code': row['salon_code'], 'qty': qty})

        cat_order = ['Salon', 'Cleaning', 'Perms', 'Tints', 'Retail']
        categories = []
        grand_total = 0
        for cat in cat_order:
            cat_items = [v for k, v in agg.items() if k[0] == cat]
            if not cat_items:
                continue
            for item in cat_items:
                t = item['total_qty']
                item['total_qty'] = int(t) if t == int(t) else t
                for s in item['salons']:
                    q = s['qty']
                    s['qty'] = int(q) if q == int(q) else q
                grand_total += item['total_qty']
            cat_total = sum(i['total_qty'] for i in cat_items)
            categories.append({
                'name': cat,
                'items': cat_items,
                'total': int(cat_total) if cat_total == int(cat_total) else cat_total,
            })

        grand_total = int(grand_total) if grand_total == int(grand_total) else grand_total
        month_label = f"{MONTHS_FULL[month-1]} {year}"
        self.render('ho/month_combined.html', user=u,
                    categories=categories, grand_total=grand_total,
                    year=year, month=month, month_label=month_label)


# ─── App ──────────────────────────────────────────────────────────────────────

def make_app():
    base = os.path.dirname(os.path.abspath(__file__))
    return tornado.web.Application(
        [
            (r'/',                              DashboardHandler),
            (r'/login',                         LoginHandler),
            (r'/logout',                        LogoutHandler),
            (r'/dashboard',                     DashboardHandler),
            (r'/salon/dashboard',               SalonDashboardHandler),
            (r'/salon/order/new',               OrderNewHandler),
            (r'/salon/order/submit',            OrderSubmitHandler),
            (r'/salon/order/preview/(\d+)',         OrderPreviewHandler),
            (r'/salon/order/preview/(\d+)/confirm', OrderConfirmHandler),
            (r'/salon/order/item/(\d+)/delete',     OrderDeleteItemHandler),
            (r'/salon/order/(\d+)/deliver',         DeliveryHandler),
            (r'/salon/order/(\d+)/revert',          OrderRevertHandler),
            (r'/salon/order/(\d+)',             OrderViewHandler),
            (r'/salon/towels',                  SalonTowelHandler),
            (r'/ho/dashboard',                  HODashboardHandler),
            (r'/ho/week/([0-9-]+)',             HOWeekHandler),
            (r'/ho/week/([0-9-]+)/pack',        HOMarkPackedHandler),
            (r'/ho/week/([0-9-]+)/packing-list',    HOPackingListHandler),
            (r'/ho/week/([0-9-]+)/print',            HOPrintViewHandler),
            (r'/ho/week/([0-9-]+)/outstanding-word', HOOutstandingDocHandler),
            (r'/ho/week/([0-9-]+)/outstanding',      HOOutstandingHandler),
            (r'/ho/week/([0-9-]+)/order-list',      HOOrderListHandler),
            (r'/ho/week/([0-9-]+)/order-list-word', HOOrderListDocHandler),
            (r'/ho/week/([0-9-]+)/supplier-order',  HOSupplierOrderSaveHandler),
            (r'/ho/week/([0-9-]+)/towels-print',    HOTowelPrintHandler),
            (r'/ho/week/([0-9-]+)/towels-word',     HOTowelWordHandler),
            (r'/ho/order/(\d+)/edit',                   HOOrderEditHandler),
            (r'/ho/order/(\d+)/fulfill-outstanding',    HOFulfillOutstandingHandler),
            (r'/ho/order/(\d+)/delete',                 HOOrderDeleteHandler),
            (r'/ho/order-item/(\d+)/mark-delivered',    HOMarkItemDeliveredHandler),
            (r'/ho/order/(\d+)',                        HOOrderViewHandler),
            (r'/ho/order/(\d+)/pack-items',             HOOrderPackItemsHandler),
            (r'/ho/products',                   AdminProductsHandler),
            (r'/ho/products/cleanup',           AdminProductCleanupHandler),
            (r'/ho/products/(\d+)/toggle',      AdminProductToggleHandler),
            (r'/ho/products/(\d+)/edit',        AdminProductEditHandler),
            (r'/ho/products/(\d+)/delete',      AdminProductDeleteHandler),
            (r'/ho/months',                      HOMonthListHandler),
            (r'/ho/month/(\d+)/(\d+)/combined',  HOMonthCombinedHandler),
            (r'/ho/month/(\d+)/(\d+)/salon/([A-Za-z0-9]+)', HOMonthSalonHandler),
            (r'/ho/month/(\d+)/(\d+)',           HOMonthHandler),
            (r'/ho/stocktake',                  HOStockTakeListHandler),
            (r'/ho/stocktake/new',              HOStockTakeNewHandler),
            (r'/ho/stocktake/(\d+)/xlsx',       HOStockTakeXlsxHandler),
            (r'/ho/stocktake/(\d+)/pdf',        HOStockTakePdfHandler),
            (r'/ho/stocktake/(\d+)',            HOStockTakeViewHandler),
            (r'/static/(.*)', tornado.web.StaticFileHandler, {'path': os.path.join(base, 'static')}),
        ],
        template_path=os.path.join(base, 'templates'),
        static_path=os.path.join(base, 'static'),
        cookie_secret=SECRET_KEY,
        login_url='/login',
        debug=True,
    )

def run_migrations():
    """Voer outomatiese DB-migrasies uit by elke opstart."""
    db = get_db()
    migrations = [
        "ALTER TABLE towel_logs ADD COLUMN next_qty REAL DEFAULT 0",
        "ALTER TABLE order_items ADD COLUMN is_custom INTEGER DEFAULT 0",
        "ALTER TABLE order_items ADD COLUMN packed INTEGER DEFAULT 0",
        "ALTER TABLE order_items ADD COLUMN delivered_qty REAL DEFAULT 0",
        "ALTER TABLE order_items ADD COLUMN delivery_note TEXT DEFAULT ''",
        "ALTER TABLE order_items ADD COLUMN note TEXT DEFAULT ''",
        "ALTER TABLE stock_orders ADD COLUMN signed_by TEXT DEFAULT ''",
        "ALTER TABLE stock_orders ADD COLUMN delivered_at TEXT",
        "ALTER TABLE stock_orders ADD COLUMN packed_at TEXT",
        "ALTER TABLE products ADD COLUMN supplier TEXT DEFAULT ''",
        "ALTER TABLE order_items ADD COLUMN packed_qty REAL DEFAULT 0",
        "ALTER TABLE order_items ADD COLUMN ho_added INTEGER DEFAULT 0",
        # Unieke indeks — verhoed duplikaat produk-name ongeag kategorie
        # Sal slegs slaag nadat duplikate verwyder is; fout word stilweg geignoreer
        "CREATE UNIQUE INDEX IF NOT EXISTS uix_products_name ON products(LOWER(TRIM(name)))",
    ]
    for stmt in migrations:
        try:
            db.execute(stmt)
            db.commit()
        except Exception:
            pass  # Kolom/indeks bestaan reeds of duplikate verhoed die skepping
    # Skep stock_takes tabelle (CREATE TABLE IF NOT EXISTS)
    db.execute("""
        CREATE TABLE IF NOT EXISTS stock_takes (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            taken_at   TEXT NOT NULL,
            taken_by   TEXT DEFAULT '',
            notes      TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    db.execute("""
        CREATE TABLE IF NOT EXISTS stock_take_items (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            stock_take_id INTEGER NOT NULL,
            product_name  TEXT NOT NULL,
            category      TEXT NOT NULL,
            quantity      REAL DEFAULT 0
        )
    """)
    # Avilan salon + gebruiker (idempotent — word oorgeslaan as dit reeds bestaan)
    db.execute("INSERT OR IGNORE INTO salons (code, name) VALUES ('AV', 'Avilan')")
    db.commit()
    avilan = db.execute("SELECT id FROM salons WHERE code='AV'").fetchone()
    if avilan and not db.execute("SELECT 1 FROM users WHERE username='reception@AV'").fetchone():
        _salt = os.urandom(16).hex()
        _h    = hashlib.pbkdf2_hmac('sha256', b'AVAV', _salt.encode(), 100000).hex()
        db.execute(
            "INSERT INTO users (username, password_hash, name, role, salon_id) VALUES (?,?,?,?,?)",
            ('reception@AV', f'{_salt}:{_h}', 'Avilan', 'salon', avilan['id'])
        )

    # ── Die Wilgers salon + gebruiker ──────────────────────────────────────────
    db.execute("INSERT OR IGNORE INTO salons (code, name) VALUES ('WG', 'Die Wilgers')")
    db.commit()
    wg = db.execute("SELECT id FROM salons WHERE code='WG'").fetchone()
    if wg and not db.execute("SELECT 1 FROM users WHERE username='reception@WG'").fetchone():
        _salt = os.urandom(16).hex()
        _h    = hashlib.pbkdf2_hmac('sha256', b'WGWG', _salt.encode(), 100000).hex()
        db.execute(
            "INSERT INTO users (username, password_hash, name, role, salon_id) VALUES (?,?,?,?,?)",
            ('reception@WG', f'{_salt}:{_h}', 'Die Wilgers', 'salon', wg['id'])
        )
        db.commit()

    # ── HO-logins: Danielia, Petro, Suan ──────────────────────────────────────
    _ho_users = [
        ('Danielia@HO', b'4444', 'Danielia'),
        ('Petro@HO',    b'5555', 'Petro'),
        ('Suan@HO',     b'6666', 'Suan'),
    ]
    for _uname, _pw, _name in _ho_users:
        if not db.execute("SELECT 1 FROM users WHERE username=?", (_uname,)).fetchone():
            _salt = os.urandom(16).hex()
            _h    = hashlib.pbkdf2_hmac('sha256', _pw, _salt.encode(), 100000).hex()
            db.execute(
                "INSERT INTO users (username, password_hash, name, role, salon_id) VALUES (?,?,?,?,?)",
                (_uname, f'{_salt}:{_h}', _name, 'ho_admin', None)
            )
    db.commit()
    # Stel alle Tints-produkte se verskaffer op HHB (idempotent)
    db.execute("UPDATE products SET supplier='HHB' WHERE category='Tints'")
    db.commit()
    # ── Migreer pack_dates na Woensdag-gebaseerde weke ──────────────────────────
    # Enige pack_date wat nie 'n Woensdag is nie (SQLite %w: 0=So,1=Ma,2=Di,3=Wo,4=Do,5=Vr,6=Sa)
    # word geskuif na die Woensdag van sy Maandag-gebaseerde kalenderweke.
    def _wed_of_week(date_str):
        d = datetime.date.fromisoformat(str(date_str)[:10])
        mon = d - datetime.timedelta(days=d.weekday())
        return (mon + datetime.timedelta(days=2)).isoformat()

    for tbl in ('stock_orders', 'supplier_orders'):
        try:
            rows = db.execute(
                f"SELECT id, pack_date FROM {tbl} WHERE strftime('%w', pack_date) != '3'"
            ).fetchall()
            for r in rows:
                db.execute(
                    f"UPDATE {tbl} SET pack_date=? WHERE id=?",
                    (_wed_of_week(r['pack_date']), r['id'])
                )
            if rows:
                db.commit()
                print(f"[Migrasie] {tbl}: {len(rows)} pack_date(s) → Woensdag")
        except Exception as e:
            print(f"[Migrasie] {tbl} pack_date migrasie fout: {e}")
    db.close()

if __name__ == '__main__':
    print(f"SVS Stock App - http://localhost:{PORT}")
    print(f"Databasis: {DB_PATH}")
    run_migrations()
    app = make_app()
    app.listen(PORT)
    import tornado.ioloop
    tornado.ioloop.IOLoop.current().start()
